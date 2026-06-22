#!/usr/bin/env python3
"""
Resume Composition Engine
Usage: python resume.py build --jd <file> --tags <tags> --template <name>
"""

import yaml
import json
import argparse
import subprocess
import os
import re
import sys
import traceback
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv

from compose import (
    DEFAULT_MAX_BULLETS,
    DEFAULT_MAX_JOBS,
    filter_skills_by_tags,
    parse_tag_list,
    rank_bullets_for_jd,
    select_experience_jobs,
)
from llm_config import (
    LLMNotConfiguredError,
    get_llm_client,
    list_providers,
    llm_chat_completion,
    resolve_llm_config,
)

BASE_FILE = "base.yaml"
LOG_FILE = "applications.json"
OUTPUT_DIR = "output"
VARIANTS_DIR = "variants"

def load_base(yaml_file: str = BASE_FILE):
    with open(yaml_file) as f:
        return yaml.safe_load(f)

def load_log():
    if Path(LOG_FILE).exists():
        with open(LOG_FILE) as f:
            return json.load(f)
    return {"applications": []}

def save_log(log):
    with open(LOG_FILE, "w") as f:
        json.dump(log, f, indent=2)

def filter_by_tags(items, tags, status_filter=["active"]):
    """Filter a list of items by tags and status."""
    if not tags:
        return [i for i in items if i.get("status", "active") in status_filter]
    return [
        i for i in items
        if i.get("status", "active") in status_filter
        and any(t in i.get("tags", []) for t in tags)
    ]

def _parse_phone(phone: str) -> str:
    """Normalize a phone number to E.164 format (+1XXXXXXXXXX)."""
    digits = re.sub(r"[^\d]", "", phone)
    if len(digits) == 10:
        return f"+1{digits}"
    if len(digits) >= 11:
        return f"+{digits}"
    return phone


def _extract_username(url: str) -> str:
    url = url.rstrip("/")
    return url.rsplit("/", 1)[-1]


def _font_for_locale(locale: str) -> str:
    """Map a locale to a rendercv font family."""
    return {
        "en": "Source Sans 3",
        "zh-CN": "Noto Sans SC",
    }.get(locale, "Source Sans 3")

def build_variant(base, tags, template, company, role, jd_text=None,
                  headline_override=None, summary_override=None, locale="en",
                  max_bullets=DEFAULT_MAX_BULLETS, max_jobs=DEFAULT_MAX_JOBS,
                  jd_keywords=None, tailored_bullets=None, jd_hard_skills=None,
                  boost_skills=False):
    """Assemble a job-specific variant from the base."""
    tags_list = parse_tag_list(tags)

    sections = {}

    summary_text = summary_override or base.get("summary")
    if summary_text:
        sections["Summary"] = [summary_text]

    exp_section = []
    for job, filtered_bullets in select_experience_jobs(
        base.get("experience", []),
        tags=tags_list,
        max_bullets=max_bullets,
        max_jobs=max_jobs,
        jd_keywords=jd_keywords,
    ):
        highlights = []
        for b in filtered_bullets:
            key = f"{job['company']}::{b['text'][:40]}"
            text = tailored_bullets.get(key, b["text"]) if tailored_bullets else b["text"]
            highlights.append(text)
        exp_section.append({
            "company": job["company"],
            "position": job["title"],
            "location": job["location"],
            "start_date": job["start"],
            "end_date": job.get("end") or "present",
            "highlights": highlights,
        })
    sections["experience"] = exp_section

    sections["skills"] = filter_skills_by_tags(
        base.get("skills", {}), tags_list,
        jd_hard_skills=jd_hard_skills,
        boost_missing=boost_skills,
    )

    sections["projects"] = [
        {
            "name": p["name"],
            "summary": p["description"],
            "highlights": [b["text"] for b in p.get("bullets", []) if b.get("status") == "active"]
        }
        for p in base.get("projects", [])
        if p.get("status") == "active"
        and (not tags_list or any(t in p.get("tags", []) for t in tags_list))
    ]

    sections["education"] = [
        {
            "institution": e["institution"],
            "area": e["degree"],
            "degree": "",
            "date": e["graduation"]
        }
        for e in base.get("education", [])
        if e.get("status") == "active"
    ]

    variant = {
        "cv": {
            "name": base["identity"]["name"],
            "email": base["identity"]["email"],
            "phone": _parse_phone(base["identity"]["phone"]),
            "location": base["identity"]["location"],
            "headline": headline_override or base["identity"].get("headline", ""),
            "photo": str(Path("..") / base["identity"]["photo"]) if base["identity"].get("photo") else None,
            "social_networks": [
                {"network": u["label"], "username": _extract_username(u["url"])}
                for u in base["identity"]["urls"]
                if u["status"] == "active"
            ],
            "sections": sections,
        },
        "design": {
            "theme": template,
            "page": {
                "size": "us-letter",
                "top_margin": "0.7in",
                "bottom_margin": "0.7in",
                "left_margin": "0.7in",
                "right_margin": "0.7in",
            },
            "colors": {
                "name": "rgb(0,79,144)",
                "headline": "rgb(0,79,144)",
                "connections": "rgb(0,79,144)",
                "section_titles": "rgb(0,79,144)",
                "links": "rgb(0,79,144)",
            },
            "typography": {
                "font_family": _font_for_locale(locale),
                "font_size": {
                    "body": "10pt",
                    "name": "30pt",
                    "headline": "10pt",
                    "connections": "10pt",
                    "section_titles": "1.4em",
                },
            },
            "header": {
                "alignment": "center",
            },
            "links": {
                "show_external_link_icon": False,
            },
        },
    }

    return variant

def write_variant(variant, slug):
    Path(VARIANTS_DIR).mkdir(exist_ok=True)
    path = f"{VARIANTS_DIR}/{slug}.yaml"
    with open(path, "w") as f:
        yaml.dump(variant, f, allow_unicode=True, sort_keys=False)
    return path

def render_variant(variant_path, slug, all_formats=False):
    """Call rendercv to render the variant. PDF-only by default."""
    output_path = str(Path(OUTPUT_DIR).resolve() / slug)
    Path(OUTPUT_DIR).mkdir(exist_ok=True)
    cmd = ["rendercv", "render", variant_path, "--output-folder", output_path]
    if not all_formats:
        cmd += ["--dont-generate-markdown", "--dont-generate-html", "--dont-generate-png"]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"rendercv error:\n{result.stderr}")
        return False
    return True


def generate_docx(variant_path: str, slug: str) -> str | None:
    """Generate a .docx file from a variant YAML. Returns output path or None on error."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        return None

    with open(variant_path) as f:
        variant = yaml.safe_load(f)

    cv = variant.get("cv", {})
    design = variant.get("design", {})
    doc = Document()

    # Page setup
    section = doc.sections[0]
    page = design.get("page", {})
    margin_str = page.get("left_margin", "0.7in")
    if margin_str.endswith("in"):
        val = float(margin_str.replace("in", ""))
        for attr in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
            setattr(section, attr, Inches(val))

    theme_color = RGBColor(0, 0x4F, 0x90)

    def _heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = theme_color
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        return p

    def _font(run, name="Calibri", size=Pt(10.5)):
        run.font.name = name
        run.font.size = size

    # Name
    name = cv.get("name", "")
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)

    # Contact line
    parts = [cv.get("email", ""), cv.get("phone", ""), cv.get("location", "")]
    contact = "  |  ".join(p for p in parts if p)
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact)
        _font(run, size=Pt(10))
        p.paragraph_format.space_after = Pt(2)

    # Headline
    headline = cv.get("headline", "")
    if headline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(headline)
        run.italic = True
        _font(run, size=Pt(10.5))
        p.paragraph_format.space_after = Pt(6)

    sections_data = cv.get("sections", {})

    # Summary
    summary_list = sections_data.get("Summary", [])
    if summary_list:
        _heading("Summary")
        for text in summary_list:
            p = doc.add_paragraph(text)
            _font(p.runs[0] if p.runs else p.add_run(), size=Pt(10.5))
            p.paragraph_format.space_after = Pt(4)

    # Experience
    exp = sections_data.get("experience", [])
    if exp:
        _heading("Experience")
        for job in exp:
            p = doc.add_paragraph()
            run = p.add_run(job.get("company", ""))
            run.bold = True
            _font(run)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(6)

            pos = job.get("position", "")
            dates = job.get("start_date", "")
            if job.get("end_date"):
                dates += f" -- {job['end_date']}"
            if pos or dates:
                p = doc.add_paragraph()
                if pos:
                    run = p.add_run(pos)
                    _font(run)
                    run.font.size = Pt(10)
                if dates:
                    run = p.add_run(f"  ({dates})")
                    run.italic = True
                    _font(run, size=Pt(10))
                p.paragraph_format.space_after = Pt(2)

            loc = job.get("location", "")
            if loc:
                p = doc.add_paragraph()
                run = p.add_run(loc)
                _font(run, size=Pt(10))
                p.paragraph_format.space_after = Pt(2)

            for hl in job.get("highlights", []):
                p = doc.add_paragraph(hl, style="List Bullet")
                for run in p.runs:
                    _font(run, size=Pt(10.5))

    # Skills
    skills = sections_data.get("skills", [])
    if skills:
        _heading("Skills")
        for sg in skills:
            label = sg.get("label", "")
            details = sg.get("details", "")
            if label or details:
                p = doc.add_paragraph()
                if label:
                    run = p.add_run(f"{label}: ")
                    run.bold = True
                    _font(run)
                if details:
                    run = p.add_run(details)
                    _font(run)
                p.paragraph_format.space_after = Pt(2)

    # Projects
    projects = sections_data.get("projects", [])
    if projects:
        _heading("Projects")
        for proj in projects:
            proj_name = proj.get("name", "")
            proj_summary = proj.get("summary", "")
            if proj_name:
                p = doc.add_paragraph()
                run = p.add_run(proj_name)
                run.bold = True
                _font(run)
                p.paragraph_format.space_after = Pt(0)
            if proj_summary:
                p = doc.add_paragraph(proj_summary)
                _font(p.runs[0] if p.runs else p.add_run(), size=Pt(10))
                p.paragraph_format.space_after = Pt(2)
            for hl in proj.get("highlights", []):
                p = doc.add_paragraph(hl, style="List Bullet")
                for run in p.runs:
                    _font(run, size=Pt(10.5))

    # Education
    edu = sections_data.get("education", [])
    if edu:
        _heading("Education")
        for e in edu:
            inst = e.get("institution", "")
            area = e.get("area", "")
            edate = e.get("date", "")
            if inst:
                p = doc.add_paragraph()
                run = p.add_run(inst)
                run.bold = True
                _font(run)
                p.paragraph_format.space_after = Pt(0)
            line = "  |  ".join(p for p in [area, edate] if p)
            if line:
                p = doc.add_paragraph(line)
                _font(p.runs[0] if p.runs else p.add_run(), size=Pt(10))
                p.paragraph_format.space_after = Pt(4)

    output_path = f"{OUTPUT_DIR}/{slug}/resume.docx"
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"DOCX written: {output_path}")
    return output_path


def log_application(slug, company, role, tags, template, jd_file):
    log = load_log()
    log["applications"].append({
        "id": slug,
        "company": company,
        "role": role,
        "date": datetime.now().isoformat()[:10],
        "tags_used": tags,
        "template": template,
        "jd_source": jd_file,
        "variant_file": f"{VARIANTS_DIR}/{slug}.yaml",
        "output_dir": f"{OUTPUT_DIR}/{slug}"
    })
    save_log(log)

def cmd_build(args):
    base = load_base(getattr(args, "yaml", BASE_FILE))

    if args.llm and not args.jd:
        print("Error: --jd is required when using --llm", file=sys.stderr)
        exit(1)
    if not args.role and not args.llm:
        print("Error: --role is required when not using --llm", file=sys.stderr)
        exit(1)

    jd_text = None
    role = args.role
    if args.jd:
        with open(args.jd) as f:
            jd_text = f.read()
        if args.llm and not role:
            role = jd_text.strip().split('\n')[0].strip()
            print(f"Extracted role from JD: {role}")

    slug = f"{args.company.lower().replace(' ','-')}-{role.lower().replace(' ','-')}-{datetime.now().strftime('%Y%m')}"

    from jd_parser import parse_jd

    parsed_jd = parse_jd(jd_text, base) if jd_text else None
    jd_keywords = parsed_jd.get("all_keywords") if parsed_jd else None

    tags = args.tags
    headline_override = None
    summary_override = None
    tailored_bullets = None
    template = args.template

    if template == "auto":
        if jd_text:
            template = select_template_auto(jd_text, base)
            print(f"Auto-selected template: {template}")
        else:
            template = "classic"
            print("No JD for auto template — using classic")

    tags_list = parse_tag_list(tags)
    top_bullets = rank_bullets_for_jd(base, tags_list, jd_keywords) if jd_text else None
    llm_provider = getattr(args, "llm_provider", None)

    if args.llm and jd_text:
        cfg = resolve_llm_config(llm_provider)
        print(f"LLM provider: {cfg['label']} ({cfg['model']} @ {cfg['base_url']})")

        print("Running LLM JD analysis...")
        tags = llm_extract_tags(jd_text, base, llm_provider=llm_provider)
        print(f"LLM suggested tags: {tags}")
        tags_list = parse_tag_list(tags)
        top_bullets = rank_bullets_for_jd(base, tags_list, jd_keywords)

        print("Generating LLM headline...")
        headline_override = llm_generate_headline(jd_text, role, llm_provider=llm_provider)
        if headline_override:
            print(f"LLM headline: {headline_override}")
        else:
            print("LLM headline failed, using base.yaml headline")

        print("Generating LLM summary...")
        summary_override = llm_generate_summary(
            jd_text, base, role, top_bullets=top_bullets, llm_provider=llm_provider,
        )
        if summary_override:
            print(f"LLM summary: {summary_override[:80]}...")
        else:
            print("LLM summary failed, using base.yaml summary")

    if getattr(args, "tailor", False) and jd_text:
        cfg = resolve_llm_config(llm_provider)
        print(f"Tailor provider: {cfg['label']} ({cfg['model']})")
        print("Tailoring bullets for JD...")
        tailored_bullets = llm_tailor_bullets(
            base, jd_text, tags, jd_keywords,
            max_bullets=args.max_bullets,
            max_jobs=args.max_jobs,
            role=role,
            llm_provider=llm_provider,
        )
        print(f"Tailored {len(tailored_bullets)} bullets")

    boost_skills = getattr(args, "boost", False)
    if boost_skills and jd_text:
        from ats import score_resume

        pre_score = score_resume(
            base, jd_text, tags=tags,
            max_bullets=args.max_bullets,
            max_jobs=args.max_jobs,
        )
        missing = pre_score["skill_match"].get("missing_skills", [])
        print(f"Boost: {len(missing)} missing hard skills detected")
        if missing:
            tailored_bullets = llm_boost_bullets(
                base, jd_text, tags, missing, tailored_bullets,
                jd_keywords,
                max_bullets=args.max_bullets,
                max_jobs=args.max_jobs,
                role=role,
                llm_provider=getattr(args, "llm_provider", None),
            )

    if not headline_override and role:
        base_headline = base["identity"].get("headline", "")
        headline_override = f"{role} | {base_headline}" if base_headline else role
        print(f"Role-based headline: {headline_override}")

    print(f"Building variant: {slug}")
    print(f"Tags: {tags}")
    print(f"Template: {template}")
    print(f"Locale: {args.locale}")
    print(f"Max bullets/job: {args.max_bullets}, Max jobs: {args.max_jobs or 'unlimited'}")

    variant = build_variant(base, tags, template, args.company, role, jd_text,
                            headline_override=headline_override,
                            summary_override=summary_override,
                            locale=args.locale,
                            max_bullets=args.max_bullets,
                            max_jobs=args.max_jobs,
                            jd_keywords=jd_keywords,
                            tailored_bullets=tailored_bullets,
                            jd_hard_skills=parsed_jd.get("hard_skills") if parsed_jd else None,
                            boost_skills=boost_skills)
    variant_path = write_variant(variant, slug)
    print(f"Variant written: {variant_path}")

    print("Rendering PDF...")
    success = render_variant(variant_path, slug, all_formats=args.all_formats)
    if success:
        print(f"Output: {OUTPUT_DIR}/{slug}/")

    log_application(slug, args.company, role, tags, template, args.jd)
    print(f"Logged to {LOG_FILE}")

    if jd_text:
        from ats import score_resume

        ats_result = score_resume(
            base, jd_text, tags=tags,
            headline=headline_override,
            summary=summary_override,
            max_bullets=args.max_bullets,
            max_jobs=args.max_jobs,
        )
        report_path = f"{OUTPUT_DIR}/{slug}/ats-report.json"
        Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
        Path(report_path).parent.mkdir(parents=True, exist_ok=True)
        with open(report_path, "w") as f:
            json.dump(ats_result, f, indent=2)
        print(f"ATS score: {ats_result['total']}/100 ({ats_result['grade']}) → {report_path}")
        if boost_skills and ats_result["total"] < 85:
            print("Tip: score below 85 — review missing skills in ats-report.json")

        if tailored_bullets:
            diff_path = f"{OUTPUT_DIR}/{slug}/bullet-diff.json"
            with open(diff_path, "w") as f:
                json.dump(tailored_bullets, f, indent=2)
            print(f"Bullet diff: {diff_path}")

    # ── Cover letter (optional) ─────────────────────────────────
    if getattr(args, "cover_letter", False):
        print("Generating cover letter...")
        cl_tags = tags or ""
        _generate_cover_letter(base, args.company, role, jd_text, cl_tags, slug,
                                yaml_file=getattr(args, "yaml", BASE_FILE))

    # ── DOCX (optional) ─────────────────────────────────────────
    if getattr(args, "docx", False):
        print("Generating DOCX...")
        generate_docx(variant_path, slug)

def cmd_tags(args):
    base = load_base()
    all_tags = set()
    for job in base.get("experience", []):
        for bullet in job.get("bullets", []):
            all_tags.update(bullet.get("tags", []))
    for cat, items in base.get("skills", {}).items():
        for item in items:
            all_tags.update(item.get("tags", []))
    print("Available tags:\n" + "\n".join(sorted(all_tags)))

def cmd_log(args):
    log = load_log()
    apps = log.get("applications", [])
    if not apps:
        print("No applications logged yet.")
        return
    for a in apps:
        print(f"\n{a['date']} — {a['company']} / {a['role']}")
        print(f"  ID:       {a['id']}")
        print(f"  Tags:     {a['tags_used']}")
        print(f"  Template: {a['template']}")
        print(f"  Output:   {a['output_dir']}")

def llm_extract_tags(jd_text, base, llm_provider=None):
    """
    Optional: call an LLM to extract relevant tags from the JD.
    Falls back to empty string if no LLM configured.
    """
    try:
        client, model, _cfg = get_llm_client(llm_provider)

        all_tags = set()
        for job in base.get("experience", []):
            for b in job.get("bullets", []):
                all_tags.update(b.get("tags", []))
        for cat, items in base.get("skills", {}).items():
            for item in items:
                all_tags.update(item.get("tags", []))

        prompt = f"""Given this job description, select the most relevant tags from the list below.
Return ONLY a comma-separated list of tags, nothing else.

Available tags: {', '.join(sorted(all_tags))}

Job description:
{jd_text[:3000]}
"""
        raw = llm_chat_completion(
            client, model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
        )
        return raw.strip()
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return ""
    except Exception as e:
        print(f"LLM error ({type(e).__name__}: {e})", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        return ""


def llm_generate_headline(jd_text: str, role: str | None = None, llm_provider=None) -> str:
    """
    Use LLM to generate a job-specific headline from the JD and target role.
    Falls back to empty string on error (caller uses base.yaml headline).
    """
    try:
        client, model, _cfg = get_llm_client(llm_provider)
        role_line = f"The target role is: {role}." if role else ""
        prompt = f"""Write a concise 1-line professional headline (10-15 words) for a resume targeting this job. {role_line} The headline MUST reflect the target role title. Include core relevant technologies. Return ONLY the headline text, nothing else.

Job description:
{jd_text[:3000]}
"""
        raw = llm_chat_completion(
            client, model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
        )
        return raw.strip().strip('"')
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return ""
    except Exception as e:
        print(f"LLM headline error ({type(e).__name__}: {e})", file=sys.stderr)
        return ""


def llm_generate_summary(jd_text: str, base: dict, role: str | None = None,
                         top_bullets: list | None = None, llm_provider=None) -> str:
    """
    Use LLM to generate a job-specific summary from the JD + top-scored bullets.
    Falls back to empty string on error (caller uses base.yaml summary).
    """
    try:
        client, model, _cfg = get_llm_client(llm_provider)

        if top_bullets:
            bullets_text = "\n".join(f"- {b['text']}" for b in top_bullets[:8])
        else:
            active_bullets = []
            for job in base.get("experience", []):
                if job.get("status") != "active":
                    continue
                for b in job.get("bullets", []):
                    if b.get("status") != "deprecated":
                        active_bullets.append(b["text"])
            bullets_text = "\n".join(f"- {b}" for b in active_bullets[:8])

        role_line = f" Target role: {role}." if role else ""

        prompt = f"""Write a professional summary for a resume targeting this job.{role_line}

Rules:
- EXACTLY 2 sentences, maximum 45 words total
- Use ONLY facts from the candidate bullets below — do not invent employers, dates, or metrics
- Include one quantified achievement if present in the source bullets
- Do NOT use filler phrases: "proven track record", "passionate", "dynamic", "results-driven"
- Return ONLY the summary text, nothing else

Candidate experience:
{bullets_text}

Job description:
{jd_text[:3000]}
"""
        raw = llm_chat_completion(
            client, model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
        )
        return raw.strip().strip('"')
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return ""
    except Exception as e:
        print(f"LLM summary error ({type(e).__name__}: {e})", file=sys.stderr)
        return ""


def llm_rewrite_cover_letter(body: str, jd_text: str, llm_provider=None) -> str:
    """Use LLM to rewrite a cover letter body to better match the JD."""
    try:
        client, model, _cfg = get_llm_client(llm_provider)
        prompt = f"""Given this cover letter template and job description, rewrite the body to better match the role. Keep the same professional tone and paragraph structure (3-4 paragraphs). Keep the opening and closing sentences intact. Return ONLY the rewritten body, nothing else.

Cover letter body:
{body}

Job description:
{jd_text[:3000]}
"""
        raw = llm_chat_completion(
            client, model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
        )
        if not raw:
            return body
        return raw.strip().strip('"')
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return body
    except Exception as e:
        print(f"LLM cover letter error ({type(e).__name__}: {e})", file=sys.stderr)
        return body


def _bullet_key(job_company: str, bullet_text: str) -> str:
    return f"{job_company}::{bullet_text[:40]}"


def _pick_bullet_text(bullet: dict) -> str:
    """Use first variant if present, else base text."""
    variants = bullet.get("variants") or []
    if variants:
        return variants[0]
    return bullet["text"]


def llm_tailor_bullets(
    base: dict,
    jd_text: str,
    tags: str | list | None,
    jd_keywords: list | None,
    max_bullets: int = DEFAULT_MAX_BULLETS,
    max_jobs: int = DEFAULT_MAX_JOBS,
    role: str | None = None,
    llm_provider: str | None = None,
) -> dict[str, str]:
    """
    Minimally rewrite selected bullets for JD alignment.
    Returns {bullet_key: tailored_text}. Skips rewrite if LLM unavailable.
    """
    try:
        client, model, _cfg = get_llm_client(llm_provider)
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return {}

    tailored: dict[str, str] = {}
    jobs = select_experience_jobs(
        base.get("experience", []),
        tags=tags,
        max_bullets=max_bullets,
        max_jobs=max_jobs,
        jd_keywords=jd_keywords,
    )

    kw_hint = ", ".join((jd_keywords or [])[:15])
    role_line = f"Target role: {role}." if role else ""

    for job, bullets in jobs:
        for bullet in bullets:
            source = _pick_bullet_text(bullet)
            key = _bullet_key(job["company"], bullet["text"])

            prompt = f"""Rewrite this resume bullet to better match the job description.

Rules:
- Keep ALL original facts — same employer, project, technologies, and metrics
- Do NOT invent numbers, tools, or achievements not in the source
- Maximum 22 words, one line
- Start with a strong action verb
- Naturally weave in relevant JD keywords if they fit: {kw_hint}
- {role_line}
- Return ONLY the rewritten bullet, nothing else

Original bullet:
{source}

Job description excerpt:
{jd_text[:2000]}
"""
            try:
                raw = llm_chat_completion(
                    client, model,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=2048,
                )
                if raw:
                    rewritten = raw.strip().strip('"')
                    if len(rewritten.split()) <= 30:
                        tailored[key] = rewritten
            except Exception as e:
                print(f"Tailor bullet skip ({type(e).__name__}): {key}", file=sys.stderr)

    return tailored


def _skill_verified_in_base(base: dict, skill: str) -> bool:
    """True if skill appears in base.yaml skills or any experience bullet."""
    skill_l = skill.lower()
    for items in base.get("skills", {}).values():
        for s in items:
            if s.get("status", "active") == "active":
                name = s.get("name", "").lower()
                if skill_l in name or name in skill_l:
                    return True
    for job in base.get("experience", []):
        for b in job.get("bullets", []):
            if b.get("status") == "deprecated":
                continue
            if skill_l in b.get("text", "").lower():
                return True
    return False


def llm_boost_bullets(
    base: dict,
    jd_text: str,
    tags: str | list | None,
    missing_skills: list[str],
    tailored: dict[str, str] | None,
    jd_keywords: list | None,
    max_bullets: int = DEFAULT_MAX_BULLETS,
    max_jobs: int = DEFAULT_MAX_JOBS,
    role: str | None = None,
    llm_provider: str | None = None,
) -> dict[str, str]:
    """
    Second-pass LLM optimization: weave verified missing hard skills into bullets.
    Only skills present in base.yaml are eligible — never fabricates experience.
    """
    verified = [s for s in missing_skills if _skill_verified_in_base(base, s)]
    if not verified:
        print("Boost: no verified missing skills to add", file=sys.stderr)
        return tailored or {}

    try:
        client, model, _cfg = get_llm_client(llm_provider)
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return tailored or {}

    boosted = dict(tailored or {})
    jobs = select_experience_jobs(
        base.get("experience", []),
        tags=tags,
        max_bullets=max_bullets,
        max_jobs=max_jobs,
        jd_keywords=jd_keywords,
    )

    missing_str = ", ".join(verified)
    role_line = f"Target role: {role}." if role else ""

    for job, bullets in jobs:
        for bullet in bullets:
            key = _bullet_key(job["company"], bullet["text"])
            current = boosted.get(key, _pick_bullet_text(bullet))
            still_missing = [s for s in verified if s.lower() not in current.lower()]
            if not still_missing:
                continue

            prompt = f"""Improve this resume bullet to naturally include these ATS keywords IF already implied by the original facts: {missing_str}

Rules:
- Do NOT invent employers, projects, metrics, or tools not in the original
- If a keyword cannot fit truthfully, return the original text unchanged
- Maximum 22 words, one line, strong action verb
- {role_line}
- Return ONLY the bullet text, nothing else

Original:
{current}

Job description excerpt:
{jd_text[:1500]}
"""
            try:
                raw = llm_chat_completion(
                    client, model,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=2048,
                )
                if raw:
                    rewritten = raw.strip().strip('"')
                    if rewritten != current and len(rewritten.split()) <= 30:
                        boosted[key] = rewritten
            except Exception as e:
                print(f"Boost skip ({type(e).__name__}): {key}", file=sys.stderr)

    added = sum(1 for k, v in boosted.items() if (tailored or {}).get(k) != v and k not in (tailored or {}))
    print(f"Boost: targeted {len(verified)} verified skills, updated bullets")
    return boosted


def select_template_auto(jd_text: str, base: dict | None = None) -> str:
    """Pick rendercv theme from JD signals."""
    from jd_parser import parse_jd

    parsed = parse_jd(jd_text, base)
    text_lower = jd_text.lower()
    seniority = parsed.get("seniority", "unknown")

    if any(w in text_lower for w in ("faang", "google", "amazon", "meta", "apple", "microsoft")):
        return "engineeringresumes"
    if seniority in ("staff", "principal", "director"):
        return "classic"
    if any(w in text_lower for w in ("startup", "early-stage", "seed", "series a")):
        return "moderncv"
    if "ats" in text_lower or "applicant tracking" in text_lower:
        return "engineeringresumes"
    return "sb2nov"


def cmd_analyze(args):
    base = load_base(getattr(args, "yaml", BASE_FILE))
    with open(args.jd) as f:
        jd_text = f.read()

    from jd_parser import parse_jd, keyword_match_report

    parsed = parse_jd(jd_text, base)
    report = keyword_match_report(parsed, base, args.tags)

    print(f"Role title:  {parsed['role_title']}")
    print(f"Seniority:   {parsed['seniority']}")
    if parsed.get("domain"):
        print(f"Domain:      {parsed['domain']}")
    print(f"\nHard skills ({len(parsed['hard_skills'])}):")
    print("  " + ", ".join(parsed["hard_skills"]) or "(none detected)")
    print(f"\nMatched in resume: {', '.join(report['matched_skills']) or '(none)'}")
    print(f"Missing from resume: {', '.join(report['missing_skills']) or '(none)'}")
    print("\nTop bullets by relevance:")
    for b in report["top_bullets"][:8]:
        print(f"  [{b['score']}] {b['job']}: {b['text'][:70]}...")

    if args.json:
        out = {**parsed, "match_report": report}
        print(json.dumps(out, indent=2))


def cmd_score(args):
    base = load_base(getattr(args, "yaml", BASE_FILE))
    with open(args.jd) as f:
        jd_text = f.read()

    from ats import score_resume

    result = score_resume(
        base, jd_text, tags=args.tags,
        max_bullets=args.max_bullets,
        max_jobs=args.max_jobs,
    )

    print(f"ATS Score: {result['total']}/100 ({result['grade']})")
    print("\nBreakdown:")
    for name, info in result["breakdown"].items():
        print(f"  {name}: {info['score']}/{info['max']}")
    print(f"\nMatched skills: {', '.join(result['skill_match']['matched_skills']) or '(none)'}")
    print(f"Missing skills: {', '.join(result['skill_match']['missing_skills']) or '(none)'}")

    if args.output:
        with open(args.output, "w") as f:
            json.dump(result, f, indent=2)
        print(f"\nReport written: {args.output}")
    elif args.json:
        print(json.dumps(result, indent=2))


def cmd_compare(args):
    base = load_base(getattr(args, "yaml", BASE_FILE))
    from ats import compare_jds

    jd_entries: list[tuple[str, str]] = []
    for path in args.jd or []:
        p = Path(path)
        if not p.exists():
            print(f"Warning: skipping missing file {path}", file=sys.stderr)
            continue
        with open(p) as f:
            jd_entries.append((p.stem, f.read()))

    if args.jds_dir:
        d = Path(args.jds_dir)
        for p in sorted(d.glob("*.txt")):
            with open(p) as f:
                jd_entries.append((p.stem, f.read()))

    if len(jd_entries) < 2:
        print("Error: provide at least 2 JD files via --jd or --jds-dir", file=sys.stderr)
        exit(1)

    if len(jd_entries) > 5:
        jd_entries = jd_entries[:5]
        print("Note: comparing first 5 JDs only", file=sys.stderr)

    result = compare_jds(
        base, jd_entries, tags=args.tags or None,
        max_bullets=args.max_bullets,
        max_jobs=args.max_jobs,
    )

    print(f"\n{'Rank':<5} {'Score':<8} {'Grade':<6} {'Role':<40} {'JD'}")
    print("-" * 90)
    for i, row in enumerate(result["rankings"], 1):
        role = (row.get("role_title") or "")[:38]
        print(f"{i:<5} {row['total']:<8} {row['grade']:<6} {role:<40} {row['label']}")
        missing = row.get("missing_skills") or []
        if missing:
            print(f"      Missing: {', '.join(missing[:6])}{'…' if len(missing) > 6 else ''}")

    print(f"\nRecommended apply first: {result['recommended']} (score {result['best_score']})")

    if args.output:
        with open(args.output, "w") as f:
            json.dump(result, f, indent=2)
        print(f"Report written: {args.output}")
    elif args.json:
        print(json.dumps(result, indent=2))


def _generate_cover_letter(base: dict, company: str, role: str | None,
                           jd_text: str | None, tags: str, slug: str, yaml_file: str = BASE_FILE):
    """Generate a cover letter .txt file in the output directory."""
    from argparse import Namespace
    cl_args = Namespace(
        yaml=yaml_file,
        company=company,
        role=role or "",
        jd=jd_text,
        tags=tags,
        llm=bool(jd_text),
        output=f"{OUTPUT_DIR}/{slug}/cover-letter-{company.lower().replace(' ','-')}.txt",
    )
    try:
        cmd_cover_letter(cl_args)
    except SystemExit:
        pass  # cmd_cover_letter calls exit(1) on error — swallow for batch mode


def cmd_cover_letter(args):
    base = load_base(args.yaml)

    if not args.company:
        print("Error: --company is required", file=sys.stderr)
        exit(1)

    jd_text = None
    if args.jd:
        with open(args.jd) as f:
            jd_text = f.read()

    role = args.role
    if args.llm and jd_text and not role:
        role = jd_text.strip().split('\n')[0].strip()

    tags_list = [t.strip() for t in args.tags.split(",") if t] if args.tags else None
    target_set = set(tags_list) if tags_list else None

    cls = base.get("cover_letters", [])
    if target_set and target_set.intersection({"ai", "fullstack"}):
        match = next((c for c in cls if c["id"] == "ai-fullstack-focused"), None)
    elif target_set and target_set.intersection({"backend", "api"}):
        match = next((c for c in cls if c["id"] == "backend-focused"), None)
    else:
        match = next((c for c in cls if c["id"] == "leadership-focused"), None)

    if not match:
        print("Error: No cover letter template found in base.yaml", file=sys.stderr)
        exit(1)

    cl_base = base.get("identity", {}).get("cover_letter_base", {})
    opening = cl_base.get("opening", "").replace("{role}", role or "{role}").replace("{company}", args.company)
    closing = cl_base.get("closing", "").replace("{role}", role or "{role}").replace("{company}", args.company)

    body = match["body"].replace("{opening}", opening).replace("{closing}", closing)
    body = body.replace("{role}", role or "{role}").replace("{company}", args.company)

    if args.llm and jd_text:
        print("Rewriting cover letter with LLM...")
        rewritten = llm_rewrite_cover_letter(body, jd_text, llm_provider=getattr(args, "llm_provider", None))
        if rewritten != body:
            print("Cover letter rewritten")
        body = rewritten

    header = f"To the Hiring Team at {args.company},\n\n"
    footer = f"\n\nBest regards,\n{base['identity']['name']}"
    full = header + body + footer

    if args.output:
        with open(args.output, "w") as f:
            f.write(full)
        print(f"Cover letter written to {args.output}")
    else:
        print("\n" + "=" * 50)
        print(f"COVER LETTER — {args.company}")
        print("=" * 50)
        print(full)


def cmd_llm_providers(args):
    """List configured LLM providers and model options."""
    active = os.environ.get("LLM_PROVIDER", "deepseek")
    print(f"Active provider (LLM_PROVIDER): {active}\n")
    for p in list_providers():
        cfg = resolve_llm_config(p["id"])
        key_set = "yes" if cfg["api_key"] else "no"
        marker = " ← active" if p["id"] == active else ""
        print(f"{p['label']} ({p['id']}){marker}")
        print(f"  API key set: {key_set}  ({p['api_key_var']})")
        print(f"  Base URL:    {cfg['base_url']}")
        print(f"  Model:       {cfg['model']}")
        print(f"  Models:      {', '.join(p['models'])}")
        print()


def main():
    load_dotenv(override=True)
    parser = argparse.ArgumentParser(description="Resume composition engine")
    subparsers = parser.add_subparsers()

    build_parser = subparsers.add_parser("build", help="Build a job-specific resume variant")
    build_parser.add_argument("--yaml", default="base.yaml", help="YAML source file (default: base.yaml)")
    build_parser.add_argument("--jd", help="Path to job description text file")
    build_parser.add_argument("--tags", default="", help="Comma-separated tags to filter by")
    build_parser.add_argument("--template", default="classic",
                              help="rendercv theme or 'auto' to pick from JD")
    build_parser.add_argument("--company", required=True, help="Company name")
    build_parser.add_argument("--role", help="Role title (extracted from JD first line if omitted with --llm)")
    build_parser.add_argument("--max-bullets", type=int, default=DEFAULT_MAX_BULLETS,
                              help="Max bullets per job (default: 4, 0=unlimited)")
    build_parser.add_argument("--max-jobs", type=int, default=DEFAULT_MAX_JOBS,
                              help="Max experience entries (default: 0=unlimited)")
    build_parser.add_argument("--locale", default="en", choices=["en", "zh-CN"],
                              help="Resume language (en or zh-CN)")
    build_parser.add_argument("--llm", action="store_true", help="Use LLM for JD analysis")
    build_parser.add_argument("--llm-provider", choices=["deepseek", "kimi", "minimax"],
                              help="LLM provider override (default: LLM_PROVIDER in .env)")
    build_parser.add_argument("--tailor", action="store_true",
                              help="LLM-rewrite selected bullets for JD (requires --jd + API key)")
    build_parser.add_argument("--boost", action="store_true",
                              help="Second LLM pass: add verified missing JD skills to bullets + skills")
    build_parser.add_argument("--all-formats", action="store_true", help="Generate HTML, Markdown, and PNG in addition to PDF")
    build_parser.add_argument("--cover-letter", action="store_true", help="Also generate a cover letter .txt file")
    build_parser.add_argument("--docx", action="store_true", help="Also generate a .docx Word document")
    build_parser.set_defaults(func=cmd_build)

    analyze_parser = subparsers.add_parser("analyze", help="Analyze a JD against base.yaml")
    analyze_parser.add_argument("--jd", required=True, help="Path to job description text file")
    analyze_parser.add_argument("--yaml", default="base.yaml", help="YAML source file")
    analyze_parser.add_argument("--tags", default="", help="Comma-separated tags filter")
    analyze_parser.add_argument("--json", action="store_true", help="Output full JSON")
    analyze_parser.set_defaults(func=cmd_analyze)

    score_parser = subparsers.add_parser("score", help="Score resume fit against a JD")
    score_parser.add_argument("--jd", required=True, help="Path to job description text file")
    score_parser.add_argument("--yaml", default="base.yaml", help="YAML source file")
    score_parser.add_argument("--tags", default="", help="Comma-separated tags filter")
    score_parser.add_argument("--max-bullets", type=int, default=DEFAULT_MAX_BULLETS)
    score_parser.add_argument("--max-jobs", type=int, default=DEFAULT_MAX_JOBS)
    score_parser.add_argument("--output", help="Write JSON report to file")
    score_parser.add_argument("--json", action="store_true", help="Print JSON to stdout")
    score_parser.set_defaults(func=cmd_score)

    compare_parser = subparsers.add_parser("compare", help="Compare resume fit against 2-5 JDs")
    compare_parser.add_argument("--jd", nargs="*", default=[], help="JD file paths (2-5)")
    compare_parser.add_argument("--jds-dir", help="Compare all .txt files in a directory")
    compare_parser.add_argument("--yaml", default="base.yaml", help="YAML source file")
    compare_parser.add_argument("--tags", default="", help="Comma-separated tags filter")
    compare_parser.add_argument("--max-bullets", type=int, default=DEFAULT_MAX_BULLETS)
    compare_parser.add_argument("--max-jobs", type=int, default=DEFAULT_MAX_JOBS)
    compare_parser.add_argument("--output", help="Write JSON report to file")
    compare_parser.add_argument("--json", action="store_true", help="Print JSON to stdout")
    compare_parser.set_defaults(func=cmd_compare)

    tags_parser = subparsers.add_parser("tags", help="List all available tags in base")
    tags_parser.set_defaults(func=cmd_tags)

    log_parser = subparsers.add_parser("log", help="Show application history")
    log_parser.set_defaults(func=cmd_log)

    cl_parser = subparsers.add_parser("cover-letter", help="Generate a cover letter from base.yaml template")
    cl_parser.add_argument("--yaml", default="base.yaml", help="YAML source file (default: base.yaml)")
    cl_parser.add_argument("--company", required=True, help="Target company name")
    cl_parser.add_argument("--role", help="Role title (extracted from JD first line if omitted with --llm)")
    cl_parser.add_argument("--jd", help="Path to job description text file")
    cl_parser.add_argument("--tags", default="", help="Comma-separated tags to select cover letter template")
    cl_parser.add_argument("--llm", action="store_true", help="Use LLM to rewrite cover letter body")
    cl_parser.add_argument("--llm-provider", choices=["deepseek", "kimi", "minimax"],
                           help="LLM provider override (default: LLM_PROVIDER in .env)")
    cl_parser.add_argument("--output", help="Output file path (default: stdout)")
    cl_parser.set_defaults(func=cmd_cover_letter)

    providers_parser = subparsers.add_parser("llm-providers", help="List LLM providers and models")
    providers_parser.set_defaults(func=cmd_llm_providers)

    args = parser.parse_args()
    if hasattr(args, "func"):
        args.func(args)
    else:
        parser.print_help()

if __name__ == "__main__":
    main()
