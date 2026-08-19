import json
import os
import sys
from contextlib import asynccontextmanager
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from sse_starlette.sse import EventSourceResponse

from src.history_db import ensure_output_dir

from .db import init_db, get_run, list_runs
from .jd_analyzer import extract_keywords, extract_text_from_pdf
from .models import (
    ResumeRunRequest,
    RunResponse,
    YamlInfo,
    JdCompareRequest,
    JdCompareResponse,
    JdPreviewRequest,
    YamlSaveRequest,
)
from .runner import start_job, stream_logs, cancel_job
from .theme_data import THEMES

REPO_ROOT = Path(__file__).resolve().parent.parent.parent
YAML_GLOBS = ["*.yaml", "*.yml"]
PROFILES_DIR = "profiles"
DEFAULT_YAML = f"{PROFILES_DIR}/career-en.yaml"
FRONTEND_DIST = REPO_ROOT / "ui" / "frontend" / "dist"


@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    ensure_output_dir()
    yield


app = FastAPI(title="Resume WebUI", lifespan=lifespan)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5300"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/health")
async def health():
    return {"status": "ok"}


def _list_yaml_files() -> list[YamlInfo]:
    from src.profiles import list_profiles

    rows = []
    for r in list_profiles(str(REPO_ROOT / PROFILES_DIR)):
        try:
            rel_path = str(Path(r["path"]).relative_to(REPO_ROOT))
        except ValueError:
            rel_path = r["path"]
        rows.append(YamlInfo(
            name=r["name"],
            path=rel_path,
            kind=r["kind"],
            market=r.get("market"),
            focus=r.get("focus"),
            source=r.get("source"),
            target_roles=r.get("target_roles"),
        ))
    return rows


def _load_yaml(yaml_file: str) -> dict:
    """Load a YAML as an effective resume dict (resolves positioning profiles)."""
    from src.profiles import load_effective

    try:
        return load_effective(str(REPO_ROOT / yaml_file))[0]
    except (FileNotFoundError, OSError):
        return {}


def _resume_text_blob(base: dict) -> str:
    parts: list[str] = []
    for job in base.get("experience", []):
        for b in job.get("bullets", []):
            if b.get("status") != "deprecated":
                parts.append(b.get("text", "").lower())
    for items in base.get("skills", {}).values():
        for s in items:
            if s.get("status", "active") == "active":
                parts.append(s.get("name", "").lower())
    return " ".join(parts)


def _jd_analysis_payload(text: str, base: dict, tags: str | list | None = None) -> dict:
    from src.jd_parser import parse_jd, keyword_match_report

    parsed = parse_jd(text, base)
    match = keyword_match_report(parsed, base, tags)
    blob = _resume_text_blob(base)
    soft_skills = parsed.get("soft_skills", [])
    domain_kw = parsed.get("domain_keywords", [])
    matched_soft = [s for s in soft_skills if s.lower() in blob]
    missing_soft = [s for s in soft_skills if s.lower() not in blob]
    matched_domain = [d for d in domain_kw if d.lower() in blob]

    return {
        "keywords": parsed.get("keywords", []),
        "hard_skills": parsed.get("hard_skills", []),
        "title_keywords": parsed.get("title_keywords", []),
        "domain_keywords": domain_kw,
        "soft_skills": soft_skills,
        "role_title": parsed.get("role_title", ""),
        "seniority": parsed.get("seniority", "unknown"),
        "domain": parsed.get("domain"),
        "matched_skills": match.get("matched_skills", []),
        "missing_skills": match.get("missing_skills", []),
        "matched_soft_skills": matched_soft,
        "missing_soft_skills": missing_soft,
        "matched_domain_keywords": matched_domain,
        "top_bullets": match.get("top_bullets", []),
    }


@app.get("/api/yamls")
async def list_yamls():
    return _list_yaml_files()


def _resolve_yaml_path(path: str) -> Path:
    """Resolve and validate a YAML path under profiles/."""
    full = (REPO_ROOT / path).resolve()
    profiles_dir = (REPO_ROOT / PROFILES_DIR).resolve()
    if not str(full).startswith(str(profiles_dir)):
        raise HTTPException(403, "Path must be under profiles/ directory")
    if full.suffix.lower() not in (".yaml", ".yml"):
        raise HTTPException(400, "File must have .yaml or .yml extension")
    return full


@app.get("/api/yaml")
async def get_yaml(path: str = DEFAULT_YAML):
    fpath = _resolve_yaml_path(path)
    if not fpath.exists():
        raise HTTPException(404, f"File not found: {path}")
    content = fpath.read_text(encoding="utf-8")
    return {"path": path, "content": content}


@app.put("/api/yaml")
async def save_yaml(data: YamlSaveRequest):
    import yaml

    # Validate YAML is parseable
    try:
        yaml.safe_load(data.content)
    except yaml.YAMLError as e:
        line = getattr(e, "problem_mark", None)
        line_num = line.line + 1 if line is not None else None
        detail = f"YAML parse error"
        if line_num:
            detail += f" on line {line_num}"
        raise HTTPException(422, detail=detail)

    fpath = _resolve_yaml_path(data.path)
    # Create backup
    backup = fpath.with_suffix(f"{fpath.suffix}~")
    if fpath.exists():
        import shutil
        shutil.copy2(str(fpath), str(backup))
    fpath.write_text(data.content, encoding="utf-8")
    return {"status": "saved", "path": data.path}


@app.get("/api/themes")
async def list_themes():
    return THEMES


@app.get("/api/fonts")
async def list_fonts():
    from src.fonts import font_choices

    return font_choices()


@app.get("/api/tags")
async def list_tags():
    base = _load_yaml(DEFAULT_YAML)
    if not base:
        return {"tags": []}

    all_tags = set()
    for job in base.get("experience", []):
        for bullet in job.get("bullets", []):
            all_tags.update(bullet.get("tags", []))
    for cat, items in base.get("skills", {}).items():
        for item in items:
            all_tags.update(item.get("tags", []))
    return {"tags": sorted(all_tags)}


@app.post("/api/jd/analyze")
async def analyze_jd(data: dict):
    text = data.get("text", "")
    if not text:
        raise HTTPException(400, "text is required")

    yaml_file = data.get("yaml_file", DEFAULT_YAML)
    base = _load_yaml(yaml_file)
    tags = data.get("tags")
    if isinstance(tags, list):
        tags = ",".join(tags) if tags else None
    return _jd_analysis_payload(text, base, tags)


@app.post("/api/jd/preview")
async def preview_jd(data: JdPreviewRequest):
    text = data.text.strip()
    if len(text) < 20:
        raise HTTPException(400, "JD text too short for preview")

    from src.jd_parser import parse_jd
    from src.compose import preview_experience_jobs

    base = _load_yaml(data.yaml_file)
    parsed = parse_jd(text, base)
    tags = ",".join(data.tags) if data.tags else None
    jobs = preview_experience_jobs(
        base.get("experience", []),
        tags=tags,
        max_bullets=data.max_bullets,
        max_jobs=data.max_jobs,
        jd_keywords=parsed.get("all_keywords"),
        priority=base.get("experience_priority"),
    )
    included_bullets = sum(
        1 for j in jobs for b in j["bullets"] if b.get("included")
    )
    excluded_bullets = sum(
        1 for j in jobs for b in j["bullets"] if not b.get("included")
    )
    return {
        "jobs": jobs,
        "jobs_included": sum(1 for j in jobs if j.get("job_included")),
        "bullets_included": included_bullets,
        "bullets_excluded": excluded_bullets,
    }


@app.post("/api/jd/upload")
async def upload_jd(file: UploadFile):
    temp = ensure_output_dir() / f".ui_temp_jd{Path(file.filename or 'file.txt').suffix}"
    content = await file.read()
    temp.write_bytes(content)
    try:
        if temp.suffix.lower() == ".pdf":
            text = extract_text_from_pdf(str(temp))
        else:
            text = temp.read_text("utf-8", errors="replace")
    finally:
        temp.unlink(missing_ok=True)

    if not text:
        raise HTTPException(400, "Could not extract text from file")

    base = _load_yaml(DEFAULT_YAML)
    payload = _jd_analysis_payload(text, base, None)
    return {"text": text[:5000], **payload}


@app.post("/api/jd/compare", response_model=JdCompareResponse)
async def compare_jds_api(data: JdCompareRequest):
    if len(data.jds) < 2:
        raise HTTPException(400, "At least 2 JDs required")
    if len(data.jds) > 5:
        raise HTTPException(400, "Maximum 5 JDs")

    from src.ats import compare_jds

    base = _load_yaml(DEFAULT_YAML)

    entries = [(item.label, item.text) for item in data.jds]
    tags = ",".join(data.tags) if data.tags else None
    result = compare_jds(
        base, entries, tags=tags,
        max_bullets=data.max_bullets,
        max_jobs=data.max_jobs,
    )
    return result


def _build_resume_cmd(args: ResumeRunRequest, jd_file: str | None) -> list[str]:
    cmd = [sys.executable, "resume.py", "build",
           "--yaml", args.yaml_file,
           "--company", args.company,
           "--template", args.theme]
    if args.role:
        cmd += ["--role", args.role]
    if args.tags:
        cmd += ["--tags", ",".join(args.tags)]
    if jd_file:
        cmd += ["--jd", jd_file]
    if not args.use_llm:
        cmd += ["--no-llm"]
    if args.llm_provider:
        cmd += ["--llm-provider", args.llm_provider]
    if args.tailor:
        cmd += ["--tailor"]
    if args.enhance:
        cmd += ["--enhance"]
    if args.boost:
        cmd += ["--boost"]
    if args.max_bullets != 4:
        cmd += ["--max-bullets", str(args.max_bullets)]
    if args.max_jobs:
        cmd += ["--max-jobs", str(args.max_jobs)]
    if args.pages != 2:
        cmd += ["--pages", str(args.pages)]
    if args.max_projects != 4:
        cmd += ["--max-projects", str(args.max_projects)]
    if args.no_projects:
        cmd += ["--no-projects"]
    if args.all_formats:
        cmd += ["--all-formats"]
    if args.locale and args.locale != "en":
        cmd += ["--locale", args.locale]
    if args.font:
        cmd += ["--font", args.font]
    if not args.cover_letter:
        cmd += ["--no-cover-letter"]
    if not args.docx:
        cmd += ["--no-docx"]
    cmd += ["--no-history"]
    return cmd


@app.post("/api/resume/run", response_model=RunResponse)
async def run_resume(args: ResumeRunRequest):
    jd_file = None
    if args.jd_text:
        jd_path = ensure_output_dir() / ".ui_temp_jd.txt"
        jd_path.write_text(args.jd_text)
        jd_file = str(jd_path)

    cmd = _build_resume_cmd(args, jd_file)
    job_id = await start_job(cmd, "resume", metadata=args.model_dump())
    return RunResponse(job_id=job_id)


@app.post("/api/resume/cancel/{job_id}")
async def cancel_run(job_id: str):
    ok = await cancel_job(job_id)
    if not ok:
        raise HTTPException(404, "Job not found or already completed")
    return {"status": "cancelled"}


@app.get("/api/log/{job_id}")
async def stream_run_log(job_id: str):
    return EventSourceResponse(stream_logs(job_id))


@app.get("/api/history")
async def get_history(
    type: str | None = None,
    status: str | None = None,
    limit: int = 50,
    offset: int = 0,
):
    rows = await list_runs(type_filter=type, status_filter=status,
                           limit=min(limit, 200), offset=offset)
    return {"runs": rows, "total": len(rows)}


@app.get("/api/history/{job_id}")
async def get_run_detail(job_id: str):
    run = await get_run(job_id)
    if not run:
        raise HTTPException(404, "Run not found")
    return run


@app.get("/api/output/{job_id}/files")
async def get_output_files(job_id: str):
    """Return list of output files for a completed run."""
    run = await get_run(job_id)
    if not run:
        raise HTTPException(404, "Run not found")
    files = run.get("output_files") or []
    # If DB has no output_files but files exist on disk, scan now
    if not files:
        output_dir = REPO_ROOT / "output"
        if output_dir.exists():
            candidates = sorted(output_dir.rglob("*"))
            # Try to match by slug from run metadata
            if run.get("company") and run.get("role"):
                from .db import scan_output_files
                files = scan_output_files(run.get("slug") or run.get("id") or "")
    return {"files": files}


async def _resolve_output_by_slug(name: str, slug: str) -> Path:
    output_dir = REPO_ROOT / "output"
    if not output_dir.exists():
        raise HTTPException(404, "No output directory")

    candidate = output_dir / slug / name
    if candidate.exists():
        return candidate

    for subdir in sorted(output_dir.iterdir(), reverse=True):
        if subdir.is_dir():
            candidate = subdir / name
            if candidate.exists():
                return candidate

    raise HTTPException(404, f"File '{name}' not found in any output directory")


@app.get("/api/output/{job_id}")
async def get_output(job_id: str, name: str | None = None):
    run = await get_run(job_id)

    if name:
        if run:
            fpath = await _resolve_output_path(job_id, name, run)
        else:
            fpath = await _resolve_output_by_slug(name, job_id)
        media_type = _mime_for_file(name)
        return FileResponse(str(fpath), media_type=media_type, filename=name)

    if run:
        # Legacy: serve first PDF
        output_path = run.get("output_path")
        if not output_path:
            output_dir = REPO_ROOT / "output"
            if output_dir.exists():
                candidates = sorted(output_dir.rglob("*.pdf"))
                if candidates:
                    output_path = str(candidates[0])
        if output_path:
            return FileResponse(output_path, media_type="application/pdf",
                                filename=Path(output_path).name)

    raise HTTPException(404, "No output found")


@app.get("/api/output/{job_id}/download")
async def download_output(job_id: str, name: str):
    """Alias for file download (used by WebUI links)."""
    return await get_output(job_id, name=name)


@app.get("/api/output/{job_id}/content")
async def get_output_content(job_id: str, name: str):
    """Return parsed JSON or raw text for an output artifact."""
    run = await get_run(job_id)
    if run:
        fpath = await _resolve_output_path(job_id, name, run)
    else:
        fpath = await _resolve_output_by_slug(name, job_id)
    text = fpath.read_text(encoding="utf-8", errors="replace")
    if fpath.suffix.lower() == ".json":
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            raise HTTPException(500, f"Invalid JSON in {name}")
    return {"name": name, "text": text}


def _infer_file_type(name: str) -> str:
    """Infer file type category from filename for the frontend."""
    lower = name.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".html") or lower.endswith(".htm"):
        return "html"
    if "cover-letter" in lower:
        return "cover-letter"
    if lower == "ats-report.json":
        return "ats-report"
    if lower == "bullet-diff.json":
        return "bullet-diff"
    if lower.endswith(".json"):
        return "json"
    if lower.endswith(".txt"):
        return "txt"
    if lower.endswith((".jpg", ".jpeg")):
        return "jpg"
    if lower.endswith(".png"):
        return "png"
    return "other"


def _mime_for_file(name: str) -> str:
    ext = Path(name).suffix.lower()
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".html": "text/html",
        ".txt": "text/plain",
        ".json": "application/json",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".typ": "text/plain",
    }.get(ext, "application/octet-stream")


@app.get("/api/outputs")
async def list_outputs():
    """List output directories from successful builds, newest first."""
    output_dir = REPO_ROOT / "output"
    if not output_dir.exists():
        return {"directories": []}

    success_slugs: set[str] = set()
    try:
        runs = await list_runs(status_filter="success", limit=200)
        for run in runs:
            for f in (run.get("output_files") or []):
                slug = f.get("slug")
                if slug:
                    success_slugs.add(slug)
    except Exception:
        pass  # DB unavailable — fall back to showing everything

    dirs = []
    for subdir in sorted(output_dir.iterdir(), reverse=True):
        if not subdir.is_dir() or subdir.name.startswith("."):
            continue
        slug = subdir.name
        if success_slugs and slug not in success_slugs:
            continue
        files = []
        for f in sorted(subdir.iterdir(), key=lambda p: p.name):
            if f.is_file() and not f.name.startswith("."):
                files.append({
                    "name": f.name,
                    "type": _infer_file_type(f.name),
                    "slug": slug,
                    "size": f.stat().st_size,
                })
        dirs.append({"slug": slug, "files": files})
    return {"directories": dirs}


@app.get("/api/outputs/view/{slug}")
async def view_output_file(slug: str, name: str):
    fpath = await _resolve_output_by_slug(name, slug)
    media_type = _mime_for_file(name)
    return FileResponse(str(fpath), media_type=media_type)


@app.get("/api/outputs/html-preview/{slug}")
async def html_preview(slug: str, name: str):
    """Convert a file to HTML for inline iframe preview. Supports DOCX, TXT, JSON."""
    fpath = await _resolve_output_by_slug(name, slug)
    ext = fpath.suffix.lower()

    if ext == ".docx":
        from docx import Document
        doc = Document(str(fpath))
        body = "".join(
            f"<p>{p.text}</p>" for p in doc.paragraphs if p.text.strip()
        )
        if not body:
            for table in doc.tables:
                body += "<table>"
                for row in table.rows:
                    body += "<tr>" + "".join(f"<td>{c.text}</td>" for c in row.cells) + "</tr>"
                body += "</table>"
        html = f"<html><body style='font-family:sans-serif;padding:16px'>{body}</body></html>"
        return HTMLResponse(html)

    if ext in (".txt", ".json", ".typ"):
        text = fpath.read_text(encoding="utf-8", errors="replace")
        html = f"<html><body style='font-family:monospace;white-space:pre-wrap;padding:16px'>{text}</body></html>"
        return HTMLResponse(html)

    from fastapi.responses import RedirectResponse
    return RedirectResponse(url=f"/api/outputs/view/{slug}?name={name}")


async def _resolve_output_path(job_id: str, name: str, run: dict) -> Path:
    output_dir = REPO_ROOT / "output"
    if not output_dir.exists():
        raise HTTPException(404, "No output directory")

    if run.get("output_files"):
        for f in run["output_files"]:
            if f["name"] == name:
                fpath = output_dir / f["slug"] / f["name"]
                if fpath.exists():
                    return fpath

    for subdir in sorted(output_dir.iterdir(), reverse=True):
        if subdir.is_dir():
            candidate = subdir / name
            if candidate.exists():
                return candidate

    raise HTTPException(404, f"File '{name}' not found")


if FRONTEND_DIST.exists():
    assets_dir = FRONTEND_DIST / "assets"
    if assets_dir.is_dir():
        app.mount("/assets", StaticFiles(directory=assets_dir), name="webui-assets")


@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    for name in ("favicon.svg", "favicon.ico"):
        candidate = FRONTEND_DIST / name
        if candidate.is_file():
            return FileResponse(candidate)
    raise HTTPException(404, "No favicon found")


if FRONTEND_DIST.exists():

    @app.get("/{full_path:path}")
    async def spa_fallback(full_path: str):
        if full_path.startswith("api/"):
            raise HTTPException(status_code=404, detail="Not found")
        candidate = FRONTEND_DIST / full_path
        if candidate.is_file():
            return FileResponse(candidate)
        return FileResponse(FRONTEND_DIST / "index.html")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("ui.backend.main:app", host="127.0.0.1", port=8000, reload=True)
