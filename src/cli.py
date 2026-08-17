#!/usr/bin/env python3
"""
Resume Composition Engine
Usage: python resume.py build --jd <file> --tags <tags> --template <name>
"""

import yaml
import json
import argparse
import subprocess
import os
import re
import shutil
import sys
import traceback
from datetime import datetime, timezone
from pathlib import Path
from dotenv import load_dotenv

import src.history_db as history_db
import src.profiles as profiles

from src.compose import (
    DEFAULT_MAX_BULLETS,
    DEFAULT_MAX_JOBS,
    bullet_key,
    filter_skills_by_tags,
    parse_tag_list,
    pick_bullet_text,
    rank_bullets_for_jd,
    select_experience_jobs,
)
from src.tailor_validation import (
    build_bullet_diff_report,
    entries_to_tailored_map,
    validate_tailor_rewrite,
)
from src.llm_config import (
    LLMNotConfiguredError,
    get_llm_client,
    list_providers,
    llm_chat_completion,
    resolve_llm_config,
)
from src.sidebar_layout import (
    compile_sidebar_outputs,
    is_sidebar_theme,
    patch_typst_for_sidebar,
)

try:
    from docx.shared import Inches, Pt, RGBColor
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

PROFILES_DIR = "profiles"
# Canonical English career source; positioning profiles layer on top of it.
BASE_FILE = f"{PROFILES_DIR}/career-en.yaml"
OUTPUT_DIR = "output"
VARIANTS_DIR = "output/variants"

def load_base(yaml_file: str = BASE_FILE):
    """Load any YAML (full source or positioning profile) as an effective base."""
    return profiles.load_effective(yaml_file)[0]

def filter_by_tags(items, tags, status_filter=["active"]):
    """Filter a list of items by tags and status."""
    if not tags:
        return [i for i in items if i.get("status", "active") in status_filter]
    return [
        i for i in items
        if i.get("status", "active") in status_filter
        and any(t in i.get("tags", []) for t in tags)
    ]

def _parse_phone(phone: str) -> str:
    """Normalize a phone number to E.164 format (+1XXXXXXXXXX)."""
    digits = re.sub(r"[^\d]", "", phone)
    if len(digits) == 10:
        return f"+1{digits}"
    if len(digits) >= 11:
        return f"+{digits}"
    return phone


def _extract_username(url: str) -> str:
    url = url.rstrip("/")
    return url.rsplit("/", 1)[-1]


def _font_for_locale(locale: str) -> str:
    """Map a locale to a rendercv font family."""
    return {
        "en": "Source Sans 3",
        "zh-CN": "Noto Sans SC",
    }.get(locale, "Source Sans 3")

# rendercv only accepts these social network names
_RENDERC_NETWORKS = {
    "LinkedIn", "GitHub", "GitLab", "IMDB", "Instagram", "ORCID",
    "Mastodon", "StackOverflow", "ResearchGate", "YouTube",
    "Google Scholar", "Telegram", "WhatsApp", "Leetcode", "X",
    "Bluesky", "Reddit",
}

# URL host fragments → valid rendercv network name
_URL_TO_NETWORK = {
    "github.com": "GitHub",
    "linkedin.com": "LinkedIn",
    "gitlab.com": "GitLab",
    "stackoverflow.com": "StackOverflow",
    "researchgate.net": "ResearchGate",
    "youtube.com": "YouTube",
    "scholar.google.com": "Google Scholar",
    "t.me": "Telegram",
    "wa.me": "WhatsApp",
    "leetcode.com": "Leetcode",
    "x.com": "X",
    "twitter.com": "X",
    "bsky.app": "Bluesky",
    "reddit.com": "Reddit",
    "imdb.com": "IMDB",
    "instagram.com": "Instagram",
    "orcid.org": "ORCID",
    "mastodon.social": "Mastodon",
}


def _resolve_network(label: str, url: str) -> str | None:
    """Map a label+url to a valid rendercv network name, or None if impossible."""
    if label in _RENDERC_NETWORKS:
        return label
    for fragment, name in _URL_TO_NETWORK.items():
        if fragment in url.lower():
            return name
    return None


def build_variant(base, tags, template, company, role, jd_text=None,
                  headline_override=None, summary_override=None, locale="en",
                  max_bullets=DEFAULT_MAX_BULLETS, max_jobs=DEFAULT_MAX_JOBS,
                  jd_keywords=None, tailored_bullets=None, jd_hard_skills=None,
                  boost_skills=False, pages=1, no_projects=False,
                  seniority=None, llm_scores=None, max_projects=4):
    """Assemble a job-specific variant from the base."""
    tags_list = parse_tag_list(tags)
    required_tags = set(tags_list) if tags_list else None

    sections = {}

    summary_text = summary_override or base.get("summary")
    if summary_text:
        sections["Summary"] = [summary_text]

    job_pairs = select_experience_jobs(
        base.get("experience", []),
        tags=tags_list,
        max_bullets=max_bullets,
        max_jobs=max_jobs,
        jd_keywords=jd_keywords,
        seniority=seniority,
        llm_scores=llm_scores,
        priority=base.get("experience_priority"),
    )

    skill_rows_preview = filter_skills_by_tags(
        base.get("skills", {}), tags_list,
        jd_hard_skills=jd_hard_skills,
        boost_missing=boost_skills,
    )
    # Score and cap projects by JD keyword overlap
    _all_projects = [
        p for p in base.get("projects", [])
        if p.get("status") == "active"
        and (not tags_list or any(t in p.get("tags", []) for t in tags_list))
    ]
    if _all_projects and jd_keywords and max_projects > 0:
        kw_set = {k.lower() for k in jd_keywords}
        def _project_score(p: dict) -> int:
            text = (p.get("description", "") + " " + " ".join(p.get("tags", []))).lower()
            return sum(1 for k in kw_set if k in text) + sum(
                1 for t in (tags_list or []) if t in p.get("tags", [])
            )
        project_list = sorted(_all_projects, key=_project_score, reverse=True)[:max_projects]
    else:
        project_list = _all_projects
    education_list = [e for e in base.get("education", []) if e.get("status") == "active"]

    page_budget_report = None
    skills_collapsed = False
    include_projects = not no_projects and bool(project_list)

    if pages > 0 and job_pairs:
        from src.page_budget import trim_jobs_to_page_budget

        job_pairs, page_budget_report = trim_jobs_to_page_budget(
            job_pairs,
            pages=pages,
            required_tags=required_tags,
            jd_keywords=jd_keywords,
            has_summary=bool(summary_text),
            skill_rows=max(len(skill_rows_preview), 1),
            project_count=len(project_list) if include_projects else 0,
            education_count=len(education_list),
        )
        skills_collapsed = page_budget_report.get("skills_collapsed", False)
        include_projects = page_budget_report.get("projects_included", include_projects)

    exp_section = []
    for job, filtered_bullets in job_pairs:
        highlights = []
        for b in filtered_bullets:
            key = bullet_key(job['company'], b['text'])
            base_text = pick_bullet_text(
                b,
                jd_keywords=jd_keywords,
                required_tags=set(tags_list) if tags_list else None,
            )
            text = tailored_bullets.get(key, base_text) if tailored_bullets else base_text
            highlights.append(text)
        exp_section.append({
            "company": job["company"],
            "position": job["title"],
            "location": job["location"],
            "start_date": job["start"],
            "end_date": job.get("end") or "present",
            "highlights": highlights,
        })
    sections["experience"] = exp_section

    skill_rows = filter_skills_by_tags(
        base.get("skills", {}), tags_list,
        jd_hard_skills=jd_hard_skills,
        boost_missing=boost_skills,
    )
    if skills_collapsed and skill_rows:
        combined = ", ".join(
            name for row in skill_rows for name in row.get("details", "").split(", ") if name
        )
        sections["skills"] = [{"label": "Skills", "details": combined}]
    else:
        sections["skills"] = skill_rows

    if include_projects:
        sections["projects"] = [
            {
                "name": p["name"],
                "summary": p["description"],
                "highlights": [b["text"] for b in p.get("bullets", []) if b.get("status") == "active"]
            }
            for p in project_list
        ]

    sections["education"] = [
        {
            "institution": e["institution"],
            "area": e["degree"],
            "degree": "",
            "date": e["graduation"]
        }
        for e in education_list
    ]

    variant = {
        "cv": {
            "name": base["identity"]["name"],
            "email": base["identity"]["email"],
            "phone": _parse_phone(base["identity"]["phone"]),
            "location": base["identity"]["location"],
            "headline": headline_override or base["identity"].get("headline", ""),
            "photo": str(Path("../..") / base["identity"]["photo"]) if base["identity"].get("photo") else None,
            "social_networks": [
                {"network": n, "username": _extract_username(u["url"])}
                for u in base["identity"]["urls"]
                if u["status"] == "active"
                and (n := _resolve_network(u["label"], u["url"])) is not None
            ],
            "sections": sections,
        },
        "design": {
            "theme": template,
        },
    }

    return variant, page_budget_report, job_pairs

def write_variant(variant, slug):
    path = Path(VARIANTS_DIR) / f"{slug}.yaml"
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w") as f:
        yaml.dump(variant, f, allow_unicode=True, sort_keys=False)
    return str(path)

def _slugify(text: str) -> str:
    """'Full-Stack Engineer' → 'Full-Stack-Engineer'"""
    return re.sub(r"[^\w.-]", "-", text).strip("-")


def _cleanup_render_byproducts(output_dir: Path):
    """Remove rendercv intermediates (typ source + copied photo) after a render."""
    for p in output_dir.glob("*.typ"):
        try:
            p.unlink(missing_ok=True)
        except OSError:
            pass
    photo = output_dir / "william-jiang.jpg"
    try:
        photo.unlink(missing_ok=True)
    except OSError:
        pass


def render_variant(variant_path, slug, all_formats=False, role=None, template=None):
    """Call rendercv to render the variant, then rename PDF with role."""
    output_path = str(Path(OUTPUT_DIR).resolve() / slug)
    Path(OUTPUT_DIR).mkdir(exist_ok=True)

    if is_sidebar_theme(template):
        return _render_variant_sidebar(
            variant_path,
            output_path,
            all_formats=all_formats,
            role=role,
        )

    cmd = ["rendercv", "render", variant_path, "--output-folder", output_path]
    if not all_formats:
        cmd += ["--dont-generate-markdown", "--dont-generate-html", "--dont-generate-png"]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"rendercv error:\n{result.stderr}")
        return False

    # rendercv names outputs after the CV's localized name (e.g. the Chinese
    # full name for zh-CN); normalize to the English filename expected below.
    pdf_files = [p for p in Path(output_path).glob("*.pdf") if p.is_file()]
    if not pdf_files:
        print(
            f"rendercv error: expected a PDF in {output_path} but none was generated",
            file=sys.stderr,
        )
        return False
    default_pdf = max(pdf_files, key=lambda p: p.stat().st_mtime)
    target_pdf = Path(output_path) / "William_Jiang_CV.pdf"
    if default_pdf != target_pdf:
        if target_pdf.exists():
            target_pdf.unlink()
        default_pdf.rename(target_pdf)
        default_pdf = target_pdf

    # Rename PDF to include role
    if role:
        custom_pdf = Path(output_path) / f"William_Jiang-{_slugify(role)}.pdf"
        if default_pdf.exists() and custom_pdf != default_pdf:
            default_pdf.rename(custom_pdf)
    _cleanup_render_byproducts(Path(output_path))
    return True


def _render_variant_sidebar(variant_path, output_path, all_formats=False, role=None):
    """Render a sidebar-theme variant: stock rendercv .typ, then our layout."""
    output_dir = Path(output_path)
    output_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        "rendercv", "render", variant_path,
        "--output-folder", output_path,
        "--dont-generate-pdf",
    ]
    if not all_formats:
        cmd += ["--dont-generate-markdown", "--dont-generate-html", "--dont-generate-png"]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"rendercv error:\n{result.stderr}")
        return False

    # rendercv only copies the photo next to the .typ when it compiles PDF/PNG;
    # with those disabled it must be copied explicitly for our own compile.
    try:
        variant = yaml.safe_load(Path(variant_path).read_text())
        photo_rel = (variant or {}).get("cv", {}).get("photo")
        if photo_rel:
            photo_src = Path(variant_path).parent / photo_rel
            shutil.copy2(photo_src, output_dir / photo_src.name)
    except (OSError, yaml.YAMLError) as e:
        print(f"sidebar render error: could not copy photo: {e}", file=sys.stderr)
        return False

    # rendercv names the .typ after the CV's localized name (e.g. the Chinese
    # full name for zh-CN); normalize it to the English filename the rest of
    # the pipeline expects, so PDF generation isn't blocked by the locale.
    typ_files = [p for p in output_dir.glob("*.typ") if p.is_file()]
    if not typ_files:
        print(
            f"rendercv error: expected a .typ file in {output_dir} but none was generated",
            file=sys.stderr,
        )
        return False
    typ_path = max(typ_files, key=lambda p: p.stat().st_mtime)
    if typ_path.name != "William_Jiang_CV.typ":
        target = output_dir / "William_Jiang_CV.typ"
        if target.exists():
            target.unlink()
        typ_path.rename(target)
        typ_path = target

    try:
        patch_typst_for_sidebar(typ_path)
        compile_sidebar_outputs(
            typ_path,
            Path(output_path),
            all_formats=all_formats,
        )
    except Exception as e:
        print(f"sidebar render error: {e}", file=sys.stderr)
        return False

    # Rename PDF to include role
    if role:
        default_pdf = Path(output_path) / "William_Jiang_CV.pdf"
        custom_pdf = Path(output_path) / f"William_Jiang-{_slugify(role)}.pdf"
        if default_pdf.exists() and custom_pdf != default_pdf:
            default_pdf.rename(custom_pdf)
    _cleanup_render_byproducts(Path(output_path))
    return True


_DOCX_NAVY = RGBColor(0x1F, 0x38, 0x64)
_DOCX_GRAY = RGBColor(0x44, 0x44, 0x44)
_DOCX_BLACK = RGBColor(0x00, 0x00, 0x00)
_DOCX_FONT = "Calibri"

_MONTHS = {
    "01": "Jan", "02": "Feb", "03": "Mar", "04": "Apr",
    "05": "May", "06": "Jun", "07": "Jul", "08": "Aug",
    "09": "Sep", "10": "Oct", "11": "Nov", "12": "Dec",
}


def _format_month_year(value) -> str:
    """'2024-10' -> 'Oct 2024'; empty -> 'Present'."""
    if not value:
        return "Present"
    parts = str(value).split("-")
    if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
        return f"{_MONTHS.get(parts[1], parts[1])} {parts[0]}"
    return str(value).title()


def _docx_setup(doc, top=0.43, bottom=0.43, left=0.49, right=0.49):
    """US Letter with Claude-reference margins; Calibri Normal style."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    normal = doc.styles["Normal"]
    normal.font.name = _DOCX_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _DOCX_GRAY


def _docx_para(doc, runs, before=None, after=None, align=None, style=None):
    """Add a paragraph from (text, size_pt, bold, color) run tuples."""
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    for text, size, bold, color in runs:
        if not text:
            continue
        r = p.add_run(text)
        r.font.name = _DOCX_FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.font.color.rgb = color
    return p


def _docx_heading(doc, text):
    return _docx_para(
        doc, [(text, 10.5, True, _DOCX_NAVY)],
        before=11, after=3,
    )


def _docx_contact_parts(cv_or_identity) -> list[str]:
    """Contact line pieces: location | phone | email | github | linkedin."""
    parts = []
    location = cv_or_identity.get("location")
    if location:
        parts.append(str(location))
    phone = cv_or_identity.get("phone")
    if phone:
        digits = re.sub(r"\D", "", str(phone))
        if len(digits) == 11 and digits.startswith("1"):
            parts.append(f"({digits[1:4]}) {digits[4:7]}-{digits[7:]}")
        else:
            parts.append(str(phone))
    email = cv_or_identity.get("email")
    if email:
        parts.append(str(email))
    urls = cv_or_identity.get("social_networks") or cv_or_identity.get("urls") or []
    seen_networks = set()
    for entry in urls:
        if isinstance(entry, dict) and entry.get("status") == "deprecated":
            continue
        label = str(entry.get("label") or entry.get("network") or "").lower()
        if "github" in label:
            if "github" in seen_networks:
                continue
            seen_networks.add("github")
            user = entry.get("username") or ""
            if not user:
                url = entry.get("url") or ""
                user = url.rstrip("/").rsplit("/", 1)[-1]
            if user:
                parts.append(f"github.com/{user}")
        elif "linkedin" in label:
            if "linkedin" in seen_networks:
                continue
            seen_networks.add("linkedin")
            user = entry.get("username") or ""
            if not user:
                url = entry.get("url") or ""
                user = url.rstrip("/").rsplit("/", 1)[-1]
            if user:
                parts.append(f"linkedin.com/in/{user}")
    return parts


def _docx_bullet(doc, text, bold_prefix=None):
    """Real Word bullet (List Bullet) with Claude-reference indentation."""
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.left_indent = Inches(0.22)
    pf.first_line_indent = Inches(-0.14)
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name = _DOCX_FONT
        r.font.size = Pt(10.5)
        r.bold = True
        r.font.color.rgb = _DOCX_BLACK
    r = p.add_run(text)
    r.font.name = _DOCX_FONT
    r.font.size = Pt(10.5)
    r.font.color.rgb = _DOCX_GRAY
    return p


def generate_docx(variant_path: str, slug: str) -> str | None:
    """Generate a Claude-style .docx resume from a variant YAML."""
    if not _HAS_DOCX:
        print("python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        return None
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        return None

    with open(variant_path) as f:
        variant = yaml.safe_load(f)

    cv = variant.get("cv", {})
    sections_data = cv.get("sections", {})
    doc = Document()
    _docx_setup(doc)

    # ── Header block (name / headline / contact) ──────────────────────────
    _docx_para(
        doc, [(str(cv.get("name", "")).upper(), 20, True, _DOCX_NAVY)],
        after=1, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    if cv.get("headline"):
        _docx_para(
            doc, [(cv["headline"], 12, False, _DOCX_BLACK)],
            after=4, align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    contact = " | ".join(_docx_contact_parts(cv))
    if contact:
        _docx_para(
            doc, [(contact, 9, False, _DOCX_GRAY)],
            after=8, align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    _docx_para(doc, [("", 10, False, _DOCX_GRAY)], after=6)

    # ── Summary ───────────────────────────────────────────────────────────
    summary_list = sections_data.get("Summary", [])
    if summary_list:
        _docx_heading(doc, "SUMMARY")
        _docx_para(
            doc, [(" ".join(str(s) for s in summary_list), 10.5, False, _DOCX_GRAY)],
            after=5,
        )

    # ── Core skills ───────────────────────────────────────────────────────
    skill_rows = sections_data.get("skills", [])
    if skill_rows:
        _docx_heading(doc, "CORE SKILLS")
        for sg in skill_rows:
            label = sg.get("label", "")
            details = sg.get("details", "")
            if not details:
                continue
            _docx_para(
                doc,
                [(f"{label}: ", 10, True, _DOCX_BLACK), (details, 10, False, _DOCX_GRAY)],
                after=2,
            )

    # ── Experience (+ earlier-career one-liners) ──────────────────────────
    exp = sections_data.get("experience", [])
    if exp:
        _docx_heading(doc, "EXPERIENCE")
        detailed, earlier = [], []
        for job in exp:
            (detailed if job.get("highlights") else earlier).append(job)
        for job in detailed:
            position = job.get("position", "")
            company = job.get("company", "")
            location = job.get("location", "")
            dates = (
                f"{_format_month_year(job.get('start_date'))} – "
                f"{_format_month_year(job.get('end_date'))}"
            )
            runs = [
                (position, 11, True, _DOCX_BLACK),
                (f" — {company}, " if company else "", 11, True, _DOCX_NAVY),
                (location, 9.5, False, _DOCX_GRAY),
                (f" | {dates}" if dates else "", 9.5, False, _DOCX_GRAY),
            ]
            _docx_para(doc, runs, before=8, after=1)
            for hl in job.get("highlights", []):
                _docx_bullet(doc, hl)
        if earlier:
            _docx_heading(doc, "EARLIER CAREER")
            for job in earlier:
                company = job.get("company", "")
                position = job.get("position", "")
                dates = (
                    f"{_format_month_year(job.get('start_date'))} – "
                    f"{_format_month_year(job.get('end_date'))}"
                )
                line = f"{company} — {position}" if position else company
                if dates:
                    line += f" | {dates}"
                _docx_para(doc, [(line, 10, False, _DOCX_GRAY)], after=1.5)

    # ── Selected projects ─────────────────────────────────────────────────
    projects = sections_data.get("projects", [])
    if projects:
        _docx_heading(doc, "SELECTED PROJECTS")
        for proj in projects:
            name = proj.get("name", "")
            summary = proj.get("summary", "")
            text = f" — {summary}" if summary else ""
            _docx_bullet(doc, text, bold_prefix=name)

    # ── Education ─────────────────────────────────────────────────────────
    edu = sections_data.get("education", [])
    if edu:
        _docx_heading(doc, "EDUCATION")
        for e in edu:
            area = e.get("area", "")
            institution = e.get("institution", "")
            edate = _format_month_year(e.get("date"))
            runs = [
                (area, 10, True, _DOCX_BLACK),
                (f" — {institution}" if institution else "", 10, False, _DOCX_GRAY),
                (f" | {edate}" if edate else "", 10, False, _DOCX_GRAY),
            ]
            _docx_para(doc, runs, after=1)

    output_path = f"{OUTPUT_DIR}/{slug}/resume.docx"
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"DOCX written: {output_path}")
    return output_path


def _write_history_from_build(slug, company, role, tags, template, args, variant_path, jd_text=None):
    """Write a build record to the shared SQLite DB."""
    from src.history_db import insert_run, scan_output_files
    now = datetime.now(timezone.utc).isoformat() if hasattr(datetime, 'timezone') else datetime.now().isoformat()
    run = {
        "id": slug,
        "type": "build",
        "status": "success",
        "yaml_file": getattr(args, "yaml", BASE_FILE),
        "company": company,
        "role": role,
        "tags": [t.strip() for t in tags.split(",") if t.strip()] if tags else [],
        "theme": getattr(args, "template", ""),
        "max_bullets": getattr(args, "max_bullets", 4),
        "max_jobs": getattr(args, "max_jobs", 0),
        "use_llm": getattr(args, "llm", False),
        "cover_letter": getattr(args, "cover_letter", False),
        "docx": getattr(args, "docx", False),
        "variant_file": variant_path,
        "jd_source": getattr(args, "jd", None) or ("output/.ui_temp_jd.txt" if jd_text else None),
        "jd_snippet": (jd_text or "")[:200],
        "created_at": now,
    }
    try:
        insert_run(run)
        files = scan_output_files(slug)
        if files:
            from src.history_db import update_run
            update_run(slug, status="success", output_files=files, output_path=str(Path(OUTPUT_DIR) / slug))
    except Exception as e:
        print(f"Warning: could not write history: {e}", file=sys.stderr)


def _update_history_ats(slug, ats_result, before_ats=None, pages=None):
    """Persist ATS score to runs.db."""
    from src.history_db import update_run, scan_output_files

    try:
        files = scan_output_files(slug)
        update_run(
            slug,
            ats_score=ats_result.get("total"),
            ats_grade=ats_result.get("grade"),
            ats_before_score=before_ats.get("total") if before_ats else None,
            pages=pages,
            output_files=files if files else None,
        )
    except Exception as e:
        print(f"Warning: could not update ATS history: {e}", file=sys.stderr)


def cmd_build(args):
    if getattr(args, "_build_attempt", 0) >= 3:
        print("Max build attempts reached.", file=sys.stderr)
        return
    args._build_attempt = getattr(args, "_build_attempt", 0) + 1

    base = load_base(getattr(args, "yaml", BASE_FILE))

    if args.llm and not args.jd:
        print("Note: --jd not provided, LLM stages will be skipped", file=sys.stderr)
        args.llm = False
    if not args.role and not args.llm:
        if args.jd:
            print("Note: --role not provided and LLM disabled, using JD first line")
    jd_text = None
    role = args.role
    if not role and not args.jd:
        default_role = (base.get("target_roles") or [None])[0]
        if default_role:
            role = default_role
            print(f"Note: --role not provided, using profile default: {role}")
        else:
            print("Error: --role is required when not using --jd", file=sys.stderr)
            exit(1)
    if args.jd:
        with open(args.jd) as f:
            jd_text = f.read()
        if args.llm and not role:
            role = jd_text.strip().split('\n')[0].strip()
            print(f"Extracted role from JD: {role}")

    slug_raw = f"{args.company.lower()}-{role.lower()}"
    slug = re.sub(r"[^a-z0-9-]", "-", slug_raw).strip("-")

    from src.jd_parser import parse_jd

    parsed_jd = parse_jd(jd_text, base) if jd_text else None
    jd_keywords = parsed_jd.get("all_keywords") if parsed_jd else None
    llm_scores: dict[str, int] = {}
    llm_provider = getattr(args, "llm_provider", None)
    rx_template = select_rx_template_auto(jd_text) if jd_text else None

    if args.llm and jd_text:
        from src.llm_pipeline import llm_parse_jd, llm_rescore_bullets

        parsed_llm = llm_parse_jd(jd_text, base, llm_provider=llm_provider)
        if parsed_llm:
            parsed_jd = parsed_llm
            jd_keywords = parsed_jd.get("all_keywords")
        llm_scores = llm_rescore_bullets(
            base, jd_text, args.tags, jd_keywords, llm_provider=llm_provider,
        )
        if llm_scores:
            print(f"LLM rescored {len(llm_scores)} top bullets (0–10)")

    tags = args.tags
    headline_override = None
    summary_override = None
    tailored_bullets = None
    bullet_diff_entries = None
    before_ats = None
    template = args.template

    if template == "auto":
        if jd_text:
            template = select_template_auto(jd_text, base)
            print(f"Auto-selected template: {template}")
        else:
            template = "classic"
            print("No JD for auto template — using classic")

    tags_list = parse_tag_list(tags)
    top_bullets = rank_bullets_for_jd(base, tags_list, jd_keywords) if jd_text else None

    if args.llm and jd_text:
        cfg = resolve_llm_config(llm_provider)
        print(f"LLM provider: {cfg['label']} ({cfg['model']} @ {cfg['base_url']})")

        print("Running LLM JD analysis...")
        tags = llm_extract_tags(jd_text, base, llm_provider=llm_provider)
        print(f"LLM suggested tags: {tags}")
        tags_list = parse_tag_list(tags)
        top_bullets = rank_bullets_for_jd(base, tags_list, jd_keywords)

        print("Generating LLM headline...")
        headline_override = llm_generate_headline(jd_text, role, llm_provider=llm_provider)
        if headline_override:
            print(f"LLM headline: {headline_override}")
        else:
            print(f"LLM headline failed, using {BASE_FILE} headline")

        print("Generating LLM summary...")
        summary_override = llm_generate_summary(
            jd_text, base, role, top_bullets=top_bullets, llm_provider=llm_provider,
        )
        if summary_override:
            print(f"LLM summary: {summary_override[:80]}...")
        else:
            print(f"LLM summary failed, using {BASE_FILE} summary")

    if getattr(args, "enhance", False) and jd_text:
        cfg = resolve_llm_config(llm_provider)
        print(f"Enhance provider: {cfg['label']} ({cfg['model']})")
        print("Enhancing experience sections for JD...")
        bullet_diff_entries = llm_enhance_experience(
            base, jd_text, tags, jd_keywords,
            max_bullets=args.max_bullets,
            max_jobs=args.max_jobs,
            role=role,
            llm_provider=llm_provider,
        )
        tailored_bullets = entries_to_tailored_map(bullet_diff_entries)
        accepted = sum(1 for e in bullet_diff_entries if e["status"] == "accepted")
        rejected = sum(1 for e in bullet_diff_entries if e["status"] == "rejected")
        unchanged = sum(1 for e in bullet_diff_entries if e["status"] == "unchanged")
        print(f"Enhanced {accepted} bullets ({unchanged} unchanged, {rejected} rejected by validation)")

    if getattr(args, "tailor", False) and jd_text:
        cfg = resolve_llm_config(llm_provider)
        print(f"Tailor provider: {cfg['label']} ({cfg['model']})")
        from src.ats import score_resume

        before_ats = score_resume(
            base, jd_text, tags=tags,
            headline=headline_override,
            summary=summary_override,
            max_bullets=args.max_bullets,
            max_jobs=args.max_jobs,
        )
        print(f"Pre-tailor ATS score: {before_ats['total']}/100 ({before_ats['grade']})")
        print("Tailoring bullets for JD...")
        bullet_diff_entries = llm_tailor_bullets(
            base, jd_text, tags, jd_keywords,
            max_bullets=args.max_bullets,
            max_jobs=args.max_jobs,
            role=role,
            llm_provider=llm_provider,
        )
        tailored_bullets = entries_to_tailored_map(bullet_diff_entries)
        accepted = sum(1 for e in bullet_diff_entries if e["status"] == "accepted")
        rejected = sum(1 for e in bullet_diff_entries if e["status"] == "rejected")
        print(f"Tailored {accepted} bullets ({rejected} rejected by validation)")

    boost_skills = getattr(args, "boost", False)
    if boost_skills and jd_text:
        from src.ats import score_resume

        if before_ats is None:
            before_ats = score_resume(
                base, jd_text, tags=tags,
                headline=headline_override,
                summary=summary_override,
                max_bullets=args.max_bullets,
                max_jobs=args.max_jobs,
            )
        pre_score = score_resume(
            base, jd_text, tags=tags,
            headline=headline_override,
            summary=summary_override,
            max_bullets=args.max_bullets,
            max_jobs=args.max_jobs,
            tailored_bullets=tailored_bullets,
        )
        missing = pre_score["skill_match"].get("missing_skills", [])
        print(f"Boost: {len(missing)} missing hard skills detected")
        if missing:
            bullet_diff_entries = llm_boost_bullets(
                base, jd_text, tags, missing,
                bullet_diff_entries,
                jd_keywords,
                max_bullets=args.max_bullets,
                max_jobs=args.max_jobs,
                role=role,
                llm_provider=getattr(args, "llm_provider", None),
            )
            tailored_bullets = entries_to_tailored_map(bullet_diff_entries)

    if not headline_override and role:
        base_headline = base["identity"].get("headline", "")
        headline_override = f"{role} | {base_headline}" if base_headline else role
        print(f"Role-based headline: {headline_override}")

    print(f"Building variant: {slug}")
    print(f"Tags: {tags}")
    print(f"Template: {template}")
    print(f"Locale: {args.locale}")
    print(f"Max bullets/job: {args.max_bullets}, Max jobs: {args.max_jobs or 'unlimited'}")
    pages = getattr(args, "pages", 1)
    if pages > 0:
        print(f"Page budget: {pages} page(s)")

    variant, page_budget_report, job_pairs = build_variant(base, tags, template, args.company, role, jd_text,
                            headline_override=headline_override,
                            summary_override=summary_override,
                            locale=args.locale,
                            max_bullets=args.max_bullets,
                            max_jobs=args.max_jobs,
                            jd_keywords=jd_keywords,
                            tailored_bullets=tailored_bullets,
                            jd_hard_skills=parsed_jd.get("hard_skills") if parsed_jd else None,
                            boost_skills=boost_skills,
                            pages=pages,
                            no_projects=getattr(args, "no_projects", False),
                            seniority=parsed_jd.get("seniority") if parsed_jd else None,
                            llm_scores=llm_scores or None,
                            max_projects=getattr(args, "max_projects", 4))
    if page_budget_report and page_budget_report.get("enabled"):
        est = page_budget_report.get("estimated_lines")
        print(f"Estimated length: ~{est} lines (budget {page_budget_report.get('budget_lines')})")
        actions = page_budget_report.get("actions") or []
        if actions:
            print(f"Page trim: {', '.join(actions)}")
    variant_path = write_variant(variant, slug)
    print(f"Variant written: {variant_path}")

    if page_budget_report and page_budget_report.get("enabled"):
        pb_path = f"{OUTPUT_DIR}/{slug}/page-budget.json"
        Path(pb_path).parent.mkdir(parents=True, exist_ok=True)
        with open(pb_path, "w") as f:
            json.dump(page_budget_report, f, indent=2)

    print("Rendering PDF...")
    success = render_variant(
        variant_path,
        slug,
        all_formats=args.all_formats,
        role=role,
        template=template,
    )
    if success:
        print(f"Output: {OUTPUT_DIR}/{slug}/")

    if not getattr(args, "no_history", False):
        _write_history_from_build(slug, args.company, role, tags, template, args, variant_path, jd_text)
        print("Logged to history DB")
    else:
        print("History logging skipped (--no-history)")

    if jd_text:
        from src.ats import score_resume

        ats_result = score_resume(
            base, jd_text, tags=tags,
            headline=headline_override,
            summary=summary_override,
            max_bullets=args.max_bullets,
            max_jobs=args.max_jobs,
            tailored_bullets=tailored_bullets,
        )
        report_path = f"{OUTPUT_DIR}/{slug}/ats-report.json"
        Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
        Path(report_path).parent.mkdir(parents=True, exist_ok=True)
        with open(report_path, "w") as f:
            json.dump(ats_result, f, indent=2)
        print(f"ATS score: {ats_result['total']}/100 ({ats_result['grade']}) → {report_path}")
        if before_ats:
            delta = round(ats_result["total"] - before_ats["total"], 1)
            sign = "+" if delta >= 0 else ""
            print(f"ATS delta (pre-tailor → final): {sign}{delta}")
        if boost_skills and ats_result["total"] < 85:
            print("Tip: score below 85 — review missing skills in ats-report.json")

        if bullet_diff_entries:
            diff_path = f"{OUTPUT_DIR}/{slug}/bullet-diff.json"
            diff_report = build_bullet_diff_report(
                bullet_diff_entries, before_ats, ats_result,
            )
            with open(diff_path, "w") as f:
                json.dump(diff_report, f, indent=2)
            print(f"Bullet diff: {diff_path}")

        if not getattr(args, "no_history", False):
            _update_history_ats(slug, ats_result, before_ats=before_ats, pages=pages if pages > 0 else None)

        from src.provenance import build_provenance_report

        prov = build_provenance_report(
            slug=slug,
            company=args.company,
            role=role,
            tags=tags,
            template=template,
            rx_template=rx_template,
            base=base,
            job_bullet_pairs=job_pairs,
            jd_keywords=jd_keywords,
            tailored_bullets=tailored_bullets,
            bullet_diff_entries=bullet_diff_entries,
            ats_result=ats_result,
            before_ats=before_ats,
            page_budget_report=page_budget_report,
            llm_scores=llm_scores or None,
            parsed_jd=parsed_jd,
        )
        prov_path = f"{OUTPUT_DIR}/{slug}/provenance.json"
        with open(prov_path, "w") as f:
            json.dump(prov, f, indent=2)
        print(f"Provenance: {prov_path}")

        target = getattr(args, "target_score", 0)
        if target and ats_result["total"] < target and args._build_attempt < 3:
            if not getattr(args, "_retried_target", False):
                print(f"Score {ats_result['total']} below target {target} — re-running with tailor+boost...")
                args._retried_target = True
                args.tailor = True
                args.boost = True
                args.llm = True
                return cmd_build(args)
            print(f"Target score {target} not reached (final: {ats_result['total']})")

    # ── Cover letter (optional) ─────────────────────────────────
    if getattr(args, "cover_letter", False):
        print("Generating cover letter...")
        cl_tags = tags or ""
        _generate_cover_letter(base, args.company, role, jd_text, cl_tags, slug,
                                yaml_file=getattr(args, "yaml", BASE_FILE),
                                use_llm=getattr(args, "llm", False))

    # ── DOCX (optional) ─────────────────────────────────────────
    if getattr(args, "docx", False):
        print("Generating DOCX...")
        generate_docx(variant_path, slug)

def cmd_tags(args):
    base = load_base()
    all_tags = set()
    for job in base.get("experience", []):
        for bullet in job.get("bullets", []):
            all_tags.update(bullet.get("tags", []))
    for cat, items in base.get("skills", {}).items():
        for item in items:
            all_tags.update(item.get("tags", []))
    print("Available tags:\n" + "\n".join(sorted(all_tags)))


def cmd_profiles(args):
    """List resume sources and positioning profiles."""
    rows = profiles.list_profiles()
    if not rows:
        print("No YAML files found in profiles/")
        return
    header = f"{'YAML':<28} {'Kind':<8} {'Market':<24} {'Focus':<48} {'Source'}"
    print(header)
    print("-" * len(header))
    for r in rows:
        market = (r.get("market") or "")[:24]
        focus = (r.get("focus") or "")[:48]
        source = Path(r["source"]).name if r.get("source") else "-"
        print(f"{r['name']:<28} {r['kind']:<8} {market:<24} {focus:<48} {source}")
    print()
    print("Usage: python resume.py build ... --yaml profiles/na-ai-engineer.yaml")
    print("Positioning profiles auto-load their source career YAML.")

def cmd_log(args):
    from src.history_db import list_runs
    runs = list_runs(limit=200)
    if not runs:
        print("No applications logged yet.")
        return
    for r in runs:
        created = (r.get("created_at") or "")[:10]
        print(f"\n{created} — {r.get('company') or '?'} / {r.get('role') or '?'}")
        print(f"  ID:       {r['id']}")
        tags_str = ", ".join(r.get("tags") or []) or "(none)"
        print(f"  Tags:     {tags_str}")
        print(f"  Template: {r.get('theme') or '-'}")
        if r.get("ats_score") is not None:
            grade = r.get("ats_grade") or "?"
            before = r.get("ats_before_score")
            delta = ""
            if before is not None:
                d = round(r["ats_score"] - before, 1)
                delta = f" (was {before}, {'+' if d >= 0 else ''}{d})"
            print(f"  ATS:      {r['ats_score']}/100 ({grade}){delta}")
        print(f"  Status:   {r.get('status', '?')}")
        dur = r.get("run_duration_seconds")
        if dur:
            print(f"  Duration: {dur:.1f}s")
        files = r.get("output_files") or []
        if files:
            print(f"  Files:    {', '.join(f['name'] for f in files)}")
        err = r.get("error_log")
        if err:
            print(f"  Error:    {err[:120]}")

def llm_extract_tags(jd_text, base, llm_provider=None):
    """
    Optional: call an LLM to extract relevant tags from the JD.
    Falls back to empty string if no LLM configured.
    """
    try:
        client, model, _cfg = get_llm_client(llm_provider)

        all_tags = set()
        for job in base.get("experience", []):
            for b in job.get("bullets", []):
                all_tags.update(b.get("tags", []))
        for cat, items in base.get("skills", {}).items():
            for item in items:
                all_tags.update(item.get("tags", []))

        prompt = f"""Given this job description, select the most relevant tags from the list below.
Return ONLY a comma-separated list of tags, nothing else.

Available tags: {', '.join(sorted(all_tags))}

Job description:
{jd_text[:3000]}
"""
        raw = llm_chat_completion(
            client, model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
        )
        return raw.strip()
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return ""
    except Exception as e:
        print(f"LLM error ({type(e).__name__}: {e})", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        return ""


def llm_generate_headline(jd_text: str, role: str | None = None, llm_provider=None) -> str:
    """
    Use LLM to generate a job-specific headline from the JD and target role.
    Falls back to empty string on error (caller uses base.yaml headline).
    """
    try:
        client, model, _cfg = get_llm_client(llm_provider)
        role_line = f"The target role is: {role}." if role else ""
        prompt = f"""Write a concise 1-line professional headline (10-15 words) for a resume targeting this job. {role_line} The headline MUST reflect the target role title. Include core relevant technologies. Return ONLY the headline text, nothing else.

Job description:
{jd_text[:3000]}
"""
        raw = llm_chat_completion(
            client, model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
        )
        return raw.strip().strip('"')
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return ""
    except Exception as e:
        print(f"LLM headline error ({type(e).__name__}: {e})", file=sys.stderr)
        return ""


def llm_generate_summary(jd_text: str, base: dict, role: str | None = None,
                         top_bullets: list | None = None, llm_provider=None) -> str:
    """
    Use LLM to generate a job-specific summary from the JD + top-scored bullets.
    Falls back to empty string on error (caller uses base.yaml summary).
    """
    try:
        client, model, _cfg = get_llm_client(llm_provider)

        if top_bullets:
            bullets_text = "\n".join(f"- {b['text']}" for b in top_bullets[:8])
        else:
            active_bullets = []
            for job in base.get("experience", []):
                if job.get("status") != "active":
                    continue
                for b in job.get("bullets", []):
                    if b.get("status") != "deprecated":
                        active_bullets.append(b["text"])
            bullets_text = "\n".join(f"- {b}" for b in active_bullets[:8])

        role_line = f" Target role: {role}." if role else ""

        prompt = f"""Write a professional summary for a resume targeting this job.{role_line}

Rules:
- EXACTLY 2 sentences, maximum 45 words total
- Use ONLY facts from the candidate bullets below — do not invent employers, dates, or metrics
- Include one quantified achievement if present in the source bullets
- Do NOT use filler phrases: "proven track record", "passionate", "dynamic", "results-driven"
- Return ONLY the summary text, nothing else

Candidate experience:
{bullets_text}

Job description:
{jd_text[:3000]}
"""
        raw = llm_chat_completion(
            client, model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
        )
        return raw.strip().strip('"')
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return ""
    except Exception as e:
        print(f"LLM summary error ({type(e).__name__}: {e})", file=sys.stderr)
        return ""


def llm_rewrite_cover_letter(body: str, jd_text: str, llm_provider=None) -> str:
    """Use LLM to rewrite a cover letter body to better match the JD."""
    try:
        client, model, _cfg = get_llm_client(llm_provider)
        prompt = f"""Given this cover letter template and job description, rewrite the body to better match the role and company.

Rules:
- Replace any references to other companies or products with the target company or generic equivalents
- Keep the same professional tone and paragraph structure (3-4 paragraphs)
- Weave in 2–3 relevant JD keywords naturally
- Return ONLY the rewritten body, nothing else

Cover letter body:
{body}

Job description:
{jd_text[:3000]}
"""
        raw = llm_chat_completion(
            client, model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
        )
        if not raw:
            return body
        return raw.strip().strip('"')
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return body
    except Exception as e:
        print(f"LLM cover letter error ({type(e).__name__}: {e})", file=sys.stderr)
        return body


def _pick_bullet_text(
    bullet: dict,
    jd_keywords: list[str] | None = None,
    tags: str | list | None = None,
) -> str:
    """Delegate to compose.pick_bullet_text for variant selection."""
    tags_list = parse_tag_list(tags)
    required = set(tags_list) if tags_list else None
    return pick_bullet_text(bullet, jd_keywords=jd_keywords, required_tags=required)


def _make_diff_entry(
    job: dict,
    bullet: dict,
    source: str,
    *,
    rewritten: str | None = None,
    status: str = "unchanged",
    rejection_reason: str | None = None,
    pass_name: str = "tailor",
) -> dict:
    original = bullet["text"]
    final = original
    if status == "accepted" and rewritten:
        final = rewritten
    elif status == "boosted" and rewritten:
        final = rewritten
    return {
        "key": bullet_key(job["company"], bullet["text"]),
        "job": job["company"],
        "title": job.get("title", ""),
        "original": original,
        "source_used": source,
        "rewritten": rewritten,
        "final": final,
        "status": status,
        "rejection_reason": rejection_reason,
        "pass": pass_name,
        "approved": status in ("accepted", "boosted"),
    }


def llm_tailor_bullets(
    base: dict,
    jd_text: str,
    tags: str | list | None,
    jd_keywords: list | None,
    max_bullets: int = DEFAULT_MAX_BULLETS,
    max_jobs: int = DEFAULT_MAX_JOBS,
    role: str | None = None,
    llm_provider: str | None = None,
) -> list[dict]:
    """
    Minimally rewrite selected bullets for JD alignment.
    Returns structured diff entries; rejected rewrites keep original text.
    """
    try:
        client, model, _cfg = get_llm_client(llm_provider)
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return []

    entries: list[dict] = []
    jobs = select_experience_jobs(
        base.get("experience", []),
        tags=tags,
        max_bullets=max_bullets,
        max_jobs=max_jobs,
        jd_keywords=jd_keywords,
        priority=base.get("experience_priority"),
    )

    kw_hint = ", ".join((jd_keywords or [])[:15])
    role_line = f"Target role: {role}." if role else ""

    for job, bullets in jobs:
        for bullet in bullets:
            source = _pick_bullet_text(bullet, jd_keywords, tags)
            entry = _make_diff_entry(job, bullet, source, pass_name="tailor")

            prompt = f"""Rewrite this resume bullet to better match the job description.

Rules:
- Keep ALL original facts — same employer, project, technologies, and metrics
- Do NOT invent numbers, tools, or achievements not in the source
- Maximum 22 words, one line
- Start with a strong action verb
- Naturally weave in relevant JD keywords if they fit: {kw_hint}
- {role_line}
- Return ONLY the rewritten bullet, nothing else

Original bullet:
{source}

Job description excerpt:
{jd_text[:2000]}
"""
            try:
                raw = llm_chat_completion(
                    client, model,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=4096,
                )
                if raw:
                    rewritten = raw.strip().strip('"')
                    if rewritten == source:
                        entry["status"] = "unchanged"
                    else:
                        ok, reason = validate_tailor_rewrite(source, rewritten)
                        if ok:
                            entry["status"] = "accepted"
                            entry["rewritten"] = rewritten
                            entry["final"] = rewritten
                            entry["approved"] = True
                        else:
                            entry["status"] = "rejected"
                            entry["rewritten"] = rewritten
                            entry["rejection_reason"] = reason
                            print(
                                f"Tailor rejected ({reason}): {entry['key']}",
                                file=sys.stderr,
                            )
            except Exception as e:
                print(
                    f"Tailor bullet skip ({type(e).__name__}): {entry['key']}",
                    file=sys.stderr,
                )

            entries.append(entry)

    return entries


def _skill_verified_in_base(base: dict, skill: str) -> bool:
    """True if skill appears in base.yaml skills or any experience bullet."""
    skill_l = skill.lower()
    for items in base.get("skills", {}).values():
        for s in items:
            if s.get("status", "active") == "active":
                name = s.get("name", "").lower()
                if skill_l in name or name in skill_l:
                    return True
    for job in base.get("experience", []):
        for b in job.get("bullets", []):
            if b.get("status") == "deprecated":
                continue
            if skill_l in b.get("text", "").lower():
                return True
    return False


def llm_boost_bullets(
    base: dict,
    jd_text: str,
    tags: str | list | None,
    missing_skills: list[str],
    entries: list[dict] | None,
    jd_keywords: list | None,
    max_bullets: int = DEFAULT_MAX_BULLETS,
    max_jobs: int = DEFAULT_MAX_JOBS,
    role: str | None = None,
    llm_provider: str | None = None,
) -> list[dict]:
    """
    Second-pass LLM optimization: weave verified missing hard skills into bullets.
    Only skills present in base.yaml are eligible — never fabricates experience.
    """
    verified = [s for s in missing_skills if _skill_verified_in_base(base, s)]
    if not verified:
        print("Boost: no verified missing skills to add", file=sys.stderr)
        return entries or []

    entry_by_key = {e["key"]: e for e in (entries or [])}

    try:
        client, model, _cfg = get_llm_client(llm_provider)
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return entries or []

    jobs = select_experience_jobs(
        base.get("experience", []),
        tags=tags,
        max_bullets=max_bullets,
        max_jobs=max_jobs,
        jd_keywords=jd_keywords,
        priority=base.get("experience_priority"),
    )

    missing_str = ", ".join(verified)
    role_line = f"Target role: {role}." if role else ""

    for job, bullets in jobs:
        for bullet in bullets:
            key = bullet_key(job["company"], bullet["text"])
            existing = entry_by_key.get(key)
            current = (
                existing["final"]
                if existing
                else _pick_bullet_text(bullet, jd_keywords, tags)
            )
            still_missing = [s for s in verified if s.lower() not in current.lower()]
            if not still_missing:
                continue

            prompt = f"""Improve this resume bullet to naturally include these ATS keywords IF already implied by the original facts: {missing_str}

Rules:
- Do NOT invent employers, projects, metrics, or tools not in the original
- If a keyword cannot fit truthfully, return the original text unchanged
- Maximum 22 words, one line, strong action verb
- {role_line}
- Return ONLY the bullet text, nothing else

Original:
{current}

Job description excerpt:
{jd_text[:1500]}
"""
            try:
                raw = llm_chat_completion(
                    client, model,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=4096,
                )
                if not raw:
                    continue
                rewritten = raw.strip().strip('"')
                if rewritten == current:
                    continue
                source_for_validation = (
                    existing["source_used"] if existing
                    else _pick_bullet_text(bullet, jd_keywords, tags)
                )
                ok, reason = validate_tailor_rewrite(source_for_validation, rewritten)
                if not ok:
                    print(f"Boost rejected ({reason}): {key}", file=sys.stderr)
                    continue
                if existing:
                    existing["rewritten"] = rewritten
                    existing["final"] = rewritten
                    existing["status"] = "boosted"
                    existing["pass"] = "boost"
                    existing["approved"] = True
                else:
                    entry_by_key[key] = _make_diff_entry(
                        job, bullet, source_for_validation,
                        rewritten=rewritten,
                        status="boosted",
                        pass_name="boost",
                    )
            except Exception as e:
                print(f"Boost skip ({type(e).__name__}): {key}", file=sys.stderr)

    print(f"Boost: targeted {len(verified)} verified skills")
    if not entry_by_key:
        return entries or []
    if entries:
        seen = {e["key"] for e in entries}
        result = [entry_by_key.get(e["key"], e) for e in entries]
        for key, entry in entry_by_key.items():
            if key not in seen:
                result.append(entry)
        return result
    return list(entry_by_key.values())


def llm_enhance_experience(
    base: dict,
    jd_text: str,
    tags: str | list | None,
    jd_keywords: list | None,
    max_bullets: int = DEFAULT_MAX_BULLETS,
    max_jobs: int = DEFAULT_MAX_JOBS,
    role: str | None = None,
    llm_provider: str | None = None,
) -> list[dict]:
    """
    LLM holistically improves each experience section: rewords bullets for impact,
    reorders by JD relevance, and suggests better job titles. Truth-first — never
    fabricates facts.
    Returns structured diff entries (same format as tailor/boost).
    """
    try:
        client, model, _cfg = get_llm_client(llm_provider)
    except LLMNotConfiguredError as e:
        print(f"LLM not configured: {e}", file=sys.stderr)
        return []

    entries: list[dict] = []
    jobs = select_experience_jobs(
        base.get("experience", []),
        tags=tags,
        max_bullets=max_bullets,
        max_jobs=max_jobs,
        jd_keywords=jd_keywords,
        priority=base.get("experience_priority"),
    )

    role_line = f"Target role: {role}." if role else ""

    for job, bullets in jobs:
        bullets_block = "\n".join(
            f"{i+1}. {_pick_bullet_text(b, jd_keywords, tags)}"
            for i, b in enumerate(bullets)
        )
        company = job["company"]
        title = job.get("title", "")

        prompt = f"""Review this resume experience section and improve it for the job description below.

Company: {company}
Current title: {title}
{role_line}

Rules:
- Reword each bullet for maximum impact: concise (≤22 words), strong action verb, quantified where possible
- Reorder bullets so the most JD-relevant ones come first
- Optionally suggest a better job title IF the current one doesn't reflect the target role well — keep it truthful to the actual role held
- Do NOT invent employers, projects, numbers, or tools not present in the source bullets
- Return ONLY a JSON object with "title" (string, same as current if unchanged) and "bullets" (array of strings in desired order), nothing else

Source bullets:
{bullets_block}

Job description excerpt:
{jd_text[:2000]}
"""
        try:
            raw = llm_chat_completion(
                client, model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=8192,
            )
            if not raw:
                for b in bullets:
                    source = _pick_bullet_text(b, jd_keywords, tags)
                    entries.append(_make_diff_entry(job, b, source, pass_name="enhance"))
                continue

            text = raw.strip()
            if text.startswith("```"):
                text = re.sub(r"^```(?:json)?\s*", "", text)
                text = re.sub(r"\s*```$", "", text)
            data = json.loads(text)

            enhanced_title = data.get("title", title)
            enhanced_bullets: list[str] = data.get("bullets", [])

            # Map enhanced bullets back to source bullets (best-effort alignment)
            source_texts = [_pick_bullet_text(b, jd_keywords, tags) for b in bullets]
            used_sources: set[int] = set()

            for eb in enhanced_bullets:
                # Find best unmatched source bullet by keyword overlap
                eb_lower = eb.lower()
                best_idx = -1
                best_overlap = 0
                for i, src in enumerate(source_texts):
                    if i in used_sources:
                        continue
                    overlap = sum(1 for w in src.lower().split() if w in eb_lower)
                    if overlap > best_overlap:
                        best_overlap = overlap
                        best_idx = i

                if best_idx >= 0:
                    used_sources.add(best_idx)
                    bullet = bullets[best_idx]
                    source = source_texts[best_idx]
                    entry = _make_diff_entry(job, bullet, source, pass_name="enhance")

                    if eb == source:
                        entry["status"] = "unchanged"
                    else:
                        ok, reason = validate_tailor_rewrite(source, eb)
                        if ok:
                            entry["status"] = "accepted"
                            entry["rewritten"] = eb
                            entry["final"] = eb
                            entry["approved"] = True
                        else:
                            entry["status"] = "rejected"
                            entry["rewritten"] = eb
                            entry["rejection_reason"] = reason
                            print(
                                f"Enhance rejected ({reason}): {entry['key']}",
                                file=sys.stderr,
                            )
                    entries.append(entry)

            # Any unmatched source bullets keep original
            for i, b in enumerate(bullets):
                if i not in used_sources:
                    source = source_texts[i]
                    entries.append(_make_diff_entry(job, b, source, pass_name="enhance"))

            if enhanced_title != title:
                print(f"Enhance title: {title} → {enhanced_title}", file=sys.stderr)

        except (json.JSONDecodeError, Exception) as e:
            print(f"Enhance skip ({type(e).__name__}) for {company}", file=sys.stderr)
            for b in bullets:
                source = _pick_bullet_text(b, jd_keywords, tags)
                entries.append(_make_diff_entry(job, b, source, pass_name="enhance"))

    return entries


def select_rx_template_auto(jd_text: str) -> str:
    """Pick rxresu.me visual template from JD signals."""
    text_lower = jd_text.lower()
    if any(w in text_lower for w in ("creative", "design", "portfolio", "visual", "brand")):
        return "bronzor"
    if any(w in text_lower for w in ("startup", "product", "founder", "pitch")):
        return "chikorita"
    return "kakuna"


def select_template_auto(jd_text: str, base: dict | None = None) -> str:
    """Pick rendercv theme from JD signals."""
    from src.jd_parser import parse_jd

    parsed = parse_jd(jd_text, base)
    text_lower = jd_text.lower()
    seniority = parsed.get("seniority", "unknown")

    if any(w in text_lower for w in ("faang", "google", "amazon", "meta", "apple", "microsoft")):
        return "engineeringresumes"
    if seniority in ("staff", "principal", "director"):
        return "classic"
    if any(w in text_lower for w in ("startup", "early-stage", "seed", "series a")):
        return "moderncv"
    if "ats" in text_lower or "applicant tracking" in text_lower:
        return "engineeringresumes"
    return "sb2nov"


def cmd_analyze(args):
    base = load_base(getattr(args, "yaml", BASE_FILE))
    with open(args.jd) as f:
        jd_text = f.read()

    from src.jd_parser import parse_jd, keyword_match_report

    parsed = parse_jd(jd_text, base)
    report = keyword_match_report(parsed, base, args.tags)

    print(f"Role title:  {parsed['role_title']}")
    print(f"Seniority:   {parsed['seniority']}")
    if parsed.get("domain"):
        print(f"Domain:      {parsed['domain']}")
    print(f"\nHard skills ({len(parsed['hard_skills'])}):")
    print("  " + ", ".join(parsed["hard_skills"]) or "(none detected)")
    print(f"\nMatched in resume: {', '.join(report['matched_skills']) or '(none)'}")
    print(f"Missing from resume: {', '.join(report['missing_skills']) or '(none)'}")
    print("\nTop bullets by relevance:")
    for b in report["top_bullets"][:8]:
        print(f"  [{b['score']}] {b['job']}: {b['text'][:70]}...")

    if args.json:
        out = {**parsed, "match_report": report}
        print(json.dumps(out, indent=2))


def cmd_score(args):
    with open(args.jd) as f:
        jd_text = f.read()

    from src.ats import score_resume, score_variant_yaml

    if getattr(args, "variant", None):
        result = score_variant_yaml(args.variant, jd_text, tags=args.tags or None)
    else:
        base = load_base(getattr(args, "yaml", BASE_FILE))
        result = score_resume(
            base, jd_text, tags=args.tags,
            max_bullets=args.max_bullets,
            max_jobs=args.max_jobs,
        )

    print(f"ATS Score: {result['total']}/100 ({result['grade']})")
    print("\nBreakdown:")
    for name, info in result["breakdown"].items():
        if "score" in info and "max" in info:
            print(f"  {name}: {info['score']}/{info['max']}")
    # Show sub-scores
    kw_sub = result["breakdown"].get("keyword_sub", {})
    if kw_sub:
        hard = kw_sub.get("hard_skills", {})
        soft = kw_sub.get("soft_skills_domain", {})
        if hard.get("pct") is not None:
            print(f"    └ hard skills: {hard['pct']}%")
        if soft.get("pct") is not None:
            print(f"    └ soft skills / domain: {soft['pct']}%")
    conc_sub = result["breakdown"].get("conciseness_sub", {})
    if conc_sub:
        print(f"    └ bullet length: {conc_sub.get('length', {}).get('pct', '—')}%")
        print(f"    └ quantified: {conc_sub.get('quantified', {}).get('pct', '—')}% ({conc_sub.get('quantified', {}).get('count', 0)} of {result.get('bullets_included', 0)})")
        print(f"    └ strong verbs: {conc_sub.get('strong_verbs', {}).get('pct', '—')}% ({conc_sub.get('strong_verbs', {}).get('count', 0)} of {result.get('bullets_included', 0)})")
    print(f"\nMatched skills: {', '.join(result['skill_match']['matched_skills']) or '(none)'}")
    if result['skill_match'].get('matched_soft_skills'):
        print(f"Matched soft skills: {', '.join(result['skill_match']['matched_soft_skills'])}")
    print(f"Missing skills: {', '.join(result['skill_match']['missing_skills']) or '(none)'}")
    if result['skill_match'].get('missing_soft_skills'):
        print(f"Missing soft skills: {', '.join(result['skill_match']['missing_soft_skills'])}")

    if args.output:
        with open(args.output, "w") as f:
            json.dump(result, f, indent=2)
        print(f"\nReport written: {args.output}")
    elif args.json:
        print(json.dumps(result, indent=2))


def cmd_interview(args):
    """Gap analysis + interview prep talking points from JD vs base.yaml."""
    base = load_base(getattr(args, "yaml", BASE_FILE))
    with open(args.jd) as f:
        jd_text = f.read()

    from src.jd_parser import parse_jd, keyword_match_report
    from src.llm_pipeline import llm_parse_jd

    parsed = parse_jd(jd_text, base)
    if args.llm:
        enriched = llm_parse_jd(jd_text, base, llm_provider=getattr(args, "llm_provider", None))
        if enriched:
            parsed = enriched

    report = keyword_match_report(parsed, base, args.tags)
    missing = report.get("missing_skills", [])
    matched = report.get("matched_skills", [])
    top = report.get("top_bullets", [])[:6]

    print(f"Interview prep — {parsed.get('role_title', 'Role')}")
    print(f"Seniority: {parsed.get('seniority', 'unknown')}\n")

    print("Strengths to lead with:")
    for s in matched[:8]:
        print(f"  ✓ {s}")
    for b in top[:4]:
        print(f"  • [{b['job']}] {b['text'][:90]}…")

    print("\nGaps to address honestly:")
    for s in missing[:10]:
        print(f"  ? {s} — prepare adjacent experience or learning narrative")

    if parsed.get("must_have_skills"):
        print("\nMust-have (from LLM parse):")
        for s in parsed["must_have_skills"]:
            mark = "✓" if s in matched else "?"
            print(f"  {mark} {s}")

    print("\nSuggested STAR stories (from top bullets):")
    for i, b in enumerate(top[:3], 1):
        print(f"  {i}. {b['job']}: expand on metrics, scope, and your specific contribution")

    if args.llm:
        try:
            client, model, _cfg = get_llm_client(getattr(args, "llm_provider", None))
            prompt = f"""Given this JD and resume gap report, list 5 likely interview questions and brief answer outlines using ONLY verified resume bullets.

Missing skills: {', '.join(missing[:12])}
Top bullets:
{chr(10).join(b['text'] for b in top)}

JD excerpt:
{jd_text[:2000]}
"""
            raw = llm_chat_completion(
                client, model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4096,
            )
            if raw:
                print("\n--- LLM interview Q&A (verify facts) ---")
                print(raw.strip())
        except LLMNotConfiguredError as e:
            print(f"\nLLM skipped: {e}", file=sys.stderr)

    if args.json:
        out = {
            "parsed_jd": parsed,
            "matched_skills": matched,
            "missing_skills": missing,
            "top_bullets": top,
        }
        print(json.dumps(out, indent=2))


def cmd_compare(args):
    base = load_base(getattr(args, "yaml", BASE_FILE))
    from src.ats import compare_jds

    jd_entries: list[tuple[str, str]] = []
    for path in args.jd or []:
        p = Path(path)
        if not p.exists():
            print(f"Warning: skipping missing file {path}", file=sys.stderr)
            continue
        with open(p) as f:
            jd_entries.append((p.stem, f.read()))

    if args.jds_dir:
        d = Path(args.jds_dir)
        for p in sorted(d.glob("*.txt")):
            with open(p) as f:
                jd_entries.append((p.stem, f.read()))

    if len(jd_entries) < 2:
        print("Error: provide at least 2 JD files via --jd or --jds-dir", file=sys.stderr)
        exit(1)

    if len(jd_entries) > 5:
        jd_entries = jd_entries[:5]
        print("Note: comparing first 5 JDs only", file=sys.stderr)

    result = compare_jds(
        base, jd_entries, tags=args.tags or None,
        max_bullets=args.max_bullets,
        max_jobs=args.max_jobs,
    )

    print(f"\n{'Rank':<5} {'Score':<8} {'Grade':<6} {'Role':<40} {'JD'}")
    print("-" * 90)
    for i, row in enumerate(result["rankings"], 1):
        role = (row.get("role_title") or "")[:38]
        print(f"{i:<5} {row['total']:<8} {row['grade']:<6} {role:<40} {row['label']}")
        missing = row.get("missing_skills") or []
        if missing:
            print(f"      Missing: {', '.join(missing[:6])}{'…' if len(missing) > 6 else ''}")

    print(f"\nRecommended apply first: {result['recommended']} (score {result['best_score']})")

    if args.output:
        with open(args.output, "w") as f:
            json.dump(result, f, indent=2)
        print(f"Report written: {args.output}")
    elif args.json:
        print(json.dumps(result, indent=2))


def _compose_cover_letter(base: dict, company: str, role: str | None,
                          jd_text: str | None, tags: str,
                          yaml_file: str = BASE_FILE,
                          use_llm: bool = False,
                          llm_provider: str | None = None) -> str | None:
    """Build the full cover-letter text (header + body + footer)."""
    cls = base.get("cover_letters", [])
    tags_list = [t.strip() for t in (tags or "").split(",") if t.strip()]
    target_set = set(tags_list) if tags_list else None
    if target_set and target_set.intersection({"ai", "fullstack"}):
        match = next((c for c in cls if c["id"] == "ai-fullstack-focused"), None)
    elif target_set and target_set.intersection({"backend", "api"}):
        match = next((c for c in cls if c["id"] == "backend-focused"), None)
    else:
        match = next((c for c in cls if c["id"] == "leadership-focused"), None)
    if not match:
        # Localized sources (e.g. base-zh-*.yaml) ship their own template ids;
        # fall back to the first cover-letter template so a letter is still
        # produced instead of being silently skipped.
        match = next(iter(cls), None)
    if not match:
        return None

    cl_base = base.get("identity", {}).get("cover_letter_base", {})
    opening = cl_base.get("opening", "").replace("{role}", role or "{role}").replace("{company}", company)
    closing = cl_base.get("closing", "").replace("{role}", role or "{role}").replace("{company}", company)
    body = match["body"].replace("{opening}", opening).replace("{closing}", closing)
    body = body.replace("{role}", role or "{role}").replace("{company}", company)

    if use_llm and jd_text:
        print("Rewriting cover letter with LLM...")
        rewritten = llm_rewrite_cover_letter(body, jd_text, llm_provider=llm_provider)
        if rewritten != body:
            print("Cover letter rewritten")
        body = rewritten

    header = f"To the Hiring Team at {company},\n\n"
    footer = f"\n\nBest regards,\n{base['identity']['name']}"
    return header + body + footer


def _generate_cover_letter(base: dict, company: str, role: str | None,
                           jd_text: str | None, tags: str, slug: str,
                           yaml_file: str = BASE_FILE, use_llm: bool = False) -> str | None:
    """Generate a styled .docx cover letter in the output directory."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    full = _compose_cover_letter(
        base, company, role, jd_text, tags, yaml_file, use_llm=use_llm,
    )
    if full is None:
        print("Cover letter: no matching template", file=sys.stderr)
        return None

    company_slug = re.sub(r"[^a-z0-9-]", "-", company.lower()).strip("-")
    out_path = Path(OUTPUT_DIR) / slug / f"cover-letter-{company_slug}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if not _HAS_DOCX:
        # Graceful fallback: plain text when python-docx is unavailable.
        txt_path = out_path.with_suffix(".txt")
        txt_path.write_text(full, encoding="utf-8")
        print(f"Cover letter written: {txt_path}")
        return str(txt_path)

    doc = Document()
    _docx_setup(doc)
    identity = base.get("identity", {})
    _docx_para(
        doc, [(str(identity.get("name", "")).upper(), 16, True, _DOCX_NAVY)],
        after=2, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    contact = " | ".join(_docx_contact_parts(identity))
    if contact:
        _docx_para(
            doc, [(contact, 9, False, _DOCX_GRAY)],
            after=14, align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for chunk in [c.strip() for c in full.split("\n\n") if c.strip()]:
        lines = [ln.strip() for ln in chunk.split("\n") if ln.strip()]
        if len(lines) == 2 and lines[0].startswith("Best regards"):
            _docx_para(doc, [(lines[0], 10.5, False, _DOCX_GRAY)], before=10, after=1)
            _docx_para(doc, [(lines[1], 10.5, True, _DOCX_BLACK)], after=0)
        else:
            _docx_para(doc, [(" ".join(lines), 10.5, False, _DOCX_GRAY)], after=8)
    doc.save(out_path)
    print(f"Cover letter written: {out_path}")
    return str(out_path)


def cmd_cover_letter(args):
    base = load_base(args.yaml)

    if not args.company:
        print("Error: --company is required", file=sys.stderr)
        exit(1)

    jd_text = None
    if args.jd:
        with open(args.jd) as f:
            jd_text = f.read()

    role = args.role
    if args.llm and jd_text and not role:
        role = jd_text.strip().split('\n')[0].strip()

    full = _compose_cover_letter(
        base, args.company, role, jd_text, args.tags, args.yaml,
        use_llm=args.llm,
        llm_provider=getattr(args, "llm_provider", None),
    )
    if full is None:
        print(f"Error: No cover letter template found in {BASE_FILE}", file=sys.stderr)
        exit(1)

    if args.output:
        out_path = Path(args.output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        with open(out_path, "w") as f:
            f.write(full)
        print(f"Cover letter written to {out_path}")
    else:
        print("\n" + "=" * 50)
        print(f"COVER LETTER — {args.company}")
        print("=" * 50)
        print(full)


def cmd_llm_providers(args):
    """List configured LLM providers and model options."""
    active = os.environ.get("LLM_PROVIDER", "deepseek")
    print(f"Active provider (LLM_PROVIDER): {active}\n")
    for p in list_providers():
        cfg = resolve_llm_config(p["id"])
        key_set = "yes" if cfg["api_key"] else "no"
        marker = " ← active" if p["id"] == active else ""
        print(f"{p['label']} ({p['id']}){marker}")
        print(f"  API key set: {key_set}  ({p['api_key_var']})")
        print(f"  Base URL:    {cfg['base_url']}")
        print(f"  Model:       {cfg['model']}")
        print(f"  Models:      {', '.join(p['models'])}")
        print()


def main():
    load_dotenv(override=True)
    history_db.init_db()
    parser = argparse.ArgumentParser(description="Resume composition engine")
    subparsers = parser.add_subparsers()

    build_parser = subparsers.add_parser("build", help="Build a job-specific resume variant")
    build_parser.add_argument("--yaml", default=BASE_FILE, help=f"YAML source file (default: {BASE_FILE})")
    build_parser.add_argument("--jd", help="Path to job description text file")
    build_parser.add_argument("--tags", default="", help="Comma-separated tags to filter by")
    build_parser.add_argument("--template", default="classic",
                              help="rendercv theme or 'auto' to pick from JD")
    build_parser.add_argument("--company", required=True, help="Company name")
    build_parser.add_argument("--role", help="Role title (extracted from JD first line if omitted with --llm)")
    build_parser.add_argument("--max-bullets", type=int, default=DEFAULT_MAX_BULLETS,
                              help="Max bullets per job (default: 4, 0=unlimited)")
    build_parser.add_argument("--max-jobs", type=int, default=DEFAULT_MAX_JOBS,
                              help="Max experience entries (default: 0=unlimited)")
    build_parser.add_argument("--pages", type=int, default=2,
                              help="Target page count for trim (default: 2, 0=disable)")
    build_parser.add_argument("--max-projects", type=int, default=4,
                              help="Max projects to include (default: 4, 0=unlimited)")
    build_parser.add_argument("--no-projects", action="store_true",
                              help="Omit projects section (also dropped by page budget when over limit)")
    build_parser.add_argument("--locale", default="en", choices=["en", "zh-CN"],
                              help="Resume language (en or zh-CN)")
    build_parser.add_argument("--llm", action=argparse.BooleanOptionalAction, default=True,
                              help="Use LLM for JD analysis (default: on, use --no-llm to disable)")
    build_parser.add_argument("--llm-provider", choices=["deepseek", "kimi", "minimax"],
                              help="LLM provider override (default: LLM_PROVIDER in .env)")
    build_parser.add_argument("--tailor", action="store_true",
                              help="LLM-rewrite selected bullets for JD (requires --jd + API key)")
    build_parser.add_argument("--enhance", action="store_true",
                              help="LLM holistically improve each experience section: reword, reorder, retitle")
    build_parser.add_argument("--boost", action="store_true",
                              help="Second LLM pass: add verified missing JD skills to bullets + skills")
    build_parser.add_argument("--target-score", type=int, default=0,
                              help="Re-run with tailor+boost if ATS score below target (e.g. 75)")
    build_parser.add_argument("--all-formats", action="store_true", help="Generate HTML, Markdown, and PNG in addition to PDF")
    build_parser.add_argument("--cover-letter", dest="cover_letter", action="store_true", default=True,
                              help="Generate a styled .docx cover letter (default: on)")
    build_parser.add_argument("--no-cover-letter", dest="cover_letter", action="store_false",
                              help="Skip cover letter generation")
    build_parser.add_argument("--docx", dest="docx", action="store_true", default=True,
                              help="Generate resume.docx (default: on)")
    build_parser.add_argument("--no-docx", dest="docx", action="store_false",
                              help="Skip DOCX generation")
    build_parser.add_argument("--no-history", action="store_true",
                              help="Skip runs.db logging (used by WebUI runner)")
    build_parser.set_defaults(func=cmd_build)

    analyze_parser = subparsers.add_parser("analyze", help=f"Analyze a JD against {BASE_FILE}")
    analyze_parser.add_argument("--jd", required=True, help="Path to job description text file")
    analyze_parser.add_argument("--yaml", default=BASE_FILE, help="YAML source file")
    analyze_parser.add_argument("--tags", default="", help="Comma-separated tags filter")
    analyze_parser.add_argument("--json", action="store_true", help="Output full JSON")
    analyze_parser.set_defaults(func=cmd_analyze)

    score_parser = subparsers.add_parser("score", help="Score resume fit against a JD")
    score_parser.add_argument("--jd", required=True, help="Path to job description text file")
    score_parser.add_argument("--yaml", default=BASE_FILE, help="YAML source file")
    score_parser.add_argument("--tags", default="", help="Comma-separated tags filter")
    score_parser.add_argument("--max-bullets", type=int, default=DEFAULT_MAX_BULLETS)
    score_parser.add_argument("--max-jobs", type=int, default=DEFAULT_MAX_JOBS)
    score_parser.add_argument("--variant", help="Score a built output/variants/*.yaml instead of composing from base")
    score_parser.add_argument("--output", help="Write JSON report to file")
    score_parser.add_argument("--json", action="store_true", help="Print JSON to stdout")
    score_parser.set_defaults(func=cmd_score)

    compare_parser = subparsers.add_parser("compare", help="Compare resume fit against 2-5 JDs")
    compare_parser.add_argument("--jd", nargs="*", default=[], help="JD file paths (2-5)")
    compare_parser.add_argument("--jds-dir", help="Compare all .txt files in a directory")
    compare_parser.add_argument("--yaml", default=BASE_FILE, help="YAML source file")
    compare_parser.add_argument("--tags", default="", help="Comma-separated tags filter")
    compare_parser.add_argument("--max-bullets", type=int, default=DEFAULT_MAX_BULLETS)
    compare_parser.add_argument("--max-jobs", type=int, default=DEFAULT_MAX_JOBS)
    compare_parser.add_argument("--output", help="Write JSON report to file")
    compare_parser.add_argument("--json", action="store_true", help="Print JSON to stdout")
    compare_parser.set_defaults(func=cmd_compare)

    interview_parser = subparsers.add_parser("interview", help="Gap analysis + interview prep from JD")
    interview_parser.add_argument("--jd", required=True, help="Path to job description text file")
    interview_parser.add_argument("--yaml", default=BASE_FILE, help="YAML source file")
    interview_parser.add_argument("--tags", default="", help="Comma-separated tags filter")
    interview_parser.add_argument("--llm", action="store_true", help="Generate LLM interview Q&A outlines")
    interview_parser.add_argument("--llm-provider", choices=["deepseek", "kimi", "minimax"],
                                  help="LLM provider override")
    interview_parser.add_argument("--json", action="store_true", help="Print JSON to stdout")
    interview_parser.set_defaults(func=cmd_interview)

    tags_parser = subparsers.add_parser("tags", help="List all available tags in base")
    tags_parser.set_defaults(func=cmd_tags)

    profiles_parser = subparsers.add_parser(
        "profiles",
        help="List resume sources and positioning profiles",
    )
    profiles_parser.set_defaults(func=cmd_profiles)

    log_parser = subparsers.add_parser("log", help="Show application history")
    log_parser.set_defaults(func=cmd_log)

    cl_parser = subparsers.add_parser("cover-letter", help=f"Generate a cover letter from {BASE_FILE} template")
    cl_parser.add_argument("--yaml", default=BASE_FILE, help=f"YAML source file (default: {BASE_FILE})")
    cl_parser.add_argument("--company", required=True, help="Target company name")
    cl_parser.add_argument("--role", help="Role title (extracted from JD first line if omitted with --llm)")
    cl_parser.add_argument("--jd", help="Path to job description text file")
    cl_parser.add_argument("--tags", default="", help="Comma-separated tags to select cover letter template")
    cl_parser.add_argument("--llm", action="store_true", help="Use LLM to rewrite cover letter body")
    cl_parser.add_argument("--llm-provider", choices=["deepseek", "kimi", "minimax"],
                           help="LLM provider override (default: LLM_PROVIDER in .env)")
    cl_parser.add_argument("--output", help="Output file path (default: stdout)")
    cl_parser.set_defaults(func=cmd_cover_letter)

    providers_parser = subparsers.add_parser("llm-providers", help="List LLM providers and models")
    providers_parser.set_defaults(func=cmd_llm_providers)

    args = parser.parse_args()
    if hasattr(args, "func"):
        args.func(args)
    else:
        parser.print_help()

if __name__ == "__main__":
    main()
